"""
production.py — наряды-допуски, инструктажи, удостоверения по профессиям
(2026-07). Отдельный модуль от миграционного учёта — своя доменная область
(охрана труда на производстве), минимально связан с остальной системой:
bot.py/webforms.py дают только пункты меню/ссылки на функции отсюда.

Реализовано "тестовым режимом" — в отдельном файле, чтобы можно было
дорабатывать/переделывать, не трогая стабильный код миграционного учёта
и табеля.
"""

from datetime import date, datetime, timedelta
import json
import logging

log = logging.getLogger("migbot.production")

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Mm, Pt
from openpyxl import Workbook
from openpyxl.styles import Alignment, Border, Font, Side
from openpyxl.worksheet.pagebreak import Break
from sqlalchemy import select
from sqlalchemy.orm import Session

from models import (
    Brigade,
    BrigadeMember,
    Certificate,
    Employee,
    Instruction,
    InstructionType,
    InternalOrder,
    OrderCategory,
    Titul,
    WorkOrder,
    WorkOrderDailyAdmission,
    WorkOrderMember,
    WorkOrderMemberChange,
    WorkLogEntry,
    WorkLogSignStatus,
    MemberChangeType,
    WorkOrderStatus,
    WorkType,
)

# Порог "скоро истекает" для удостоверений — по аналогии с обязательствами
# миграционного учёта. [Предполагаю] 30 дней — не согласовано явно, разумный
# дефолт, легко поменять одной константой.
CERTIFICATE_EXPIRY_WARNING_DAYS = 30


# ================= Бригады =================

def create_brigade(session: Session, name: str, member_employee_ids: list[str]) -> Brigade:
    brigade = Brigade(name=name)
    session.add(brigade)
    session.flush()
    for employee_id in member_employee_ids:
        session.add(BrigadeMember(brigade_id=brigade.id, employee_id=employee_id))
    session.commit()
    return brigade


def get_brigades(session: Session) -> list[Brigade]:
    return session.query(Brigade).order_by(Brigade.name).all()


def get_brigade_member_ids(session: Session, brigade_id: str) -> list[str]:
    return [
        m.employee_id for m in
        session.query(BrigadeMember).filter_by(brigade_id=brigade_id).all()
    ]


def update_brigade_members(session: Session, brigade_id: str, member_employee_ids: list[str]) -> bool:
    """Заменяет состав целиком (не история — просто текущий список, см. docstring
    Brigade в models.py)."""
    brigade = session.get(Brigade, brigade_id)
    if brigade is None:
        return False
    session.query(BrigadeMember).filter_by(brigade_id=brigade_id).delete()
    for employee_id in member_employee_ids:
        session.add(BrigadeMember(brigade_id=brigade_id, employee_id=employee_id))
    session.commit()
    return True


def update_brigade(session: Session, brigade_id: str, name: str,
                   member_employee_ids: list[str]) -> bool:
    """Переименование + замена состава одной операцией (одна форма редактирования
    в вебе). Пустое/пробельное имя игнорируется — остаётся прежнее, чтобы случайно
    не обнулить название."""
    brigade = session.get(Brigade, brigade_id)
    if brigade is None:
        return False
    if name and name.strip():
        brigade.name = name.strip()
    session.query(BrigadeMember).filter_by(brigade_id=brigade_id).delete()
    for employee_id in member_employee_ids:
        session.add(BrigadeMember(brigade_id=brigade_id, employee_id=employee_id))
    session.add(brigade)
    session.commit()
    return True


def delete_brigade(session: Session, brigade_id: str) -> bool:
    brigade = session.get(Brigade, brigade_id)
    if brigade is None:
        return False
    # Явно удаляем членов перед бригадой — не полагаемся на каскад в модели
    # (если он есть, лишним не будет; если нет — не остаётся сирот BrigadeMember).
    session.query(BrigadeMember).filter_by(brigade_id=brigade_id).delete()
    session.delete(brigade)
    session.commit()
    return True


# ================= Титулы (объекты работ) =================
# Плоский справочник для поля «Место выполнения работ» наряда-допуска. По образцу
# бригад: своя сущность + страница ведения (/production/tituly), в наряде select
# наполняет WorkOrder.location. Сам титул на наряде НЕ хранится — location остаётся
# строкой, поэтому разовый объект не из справочника можно вписать вручную.
# code — шифр (15.21), name — наименование (Лаборатория).

def create_titul(session: Session, code: str, name: str) -> Titul:
    titul = Titul(code=code.strip(), name=name.strip())
    session.add(titul)
    session.commit()
    return titul


def get_tituly(session: Session) -> list[Titul]:
    return (
        session.query(Titul)
        .order_by(Titul.code)
        .all()
    )


def update_titul(session: Session, titul_id: str, code: str, name: str) -> bool:
    """Пустое/пробельное значение игнорируется — прежнее не обнуляется (как в
    update_brigade)."""
    titul = session.get(Titul, titul_id)
    if titul is None:
        return False
    if code and code.strip():
        titul.code = code.strip()
    if name and name.strip():
        titul.name = name.strip()
    session.add(titul)
    session.commit()
    return True


def delete_titul(session: Session, titul_id: str) -> bool:
    titul = session.get(Titul, titul_id)
    if titul is None:
        return False
    session.delete(titul)
    session.commit()
    return True


# ================= Наряды-допуски =================

def create_work_order(
    session: Session, number: str, work_description: str, location: str,
    responsible_supervisor_id: str, responsible_executor_id: str, issued_by: str,
    valid_from: date, valid_to: date, member_employee_ids: list[str],
    subdivision: str | None = None, materials: str | None = None, tools: str | None = None,
    equipment: str | None = None, special_machinery: str | None = None,
    technological_card_ref: str | None = None, safety_systems: str | None = None,
    special_conditions: str | None = None, work_type_id: str | None = None,
    titul_id: str | None = None,
) -> WorkOrder:
    order = WorkOrder(
        number=number,
        subdivision=subdivision,
        work_description=work_description,
        location=location,
        responsible_supervisor_id=responsible_supervisor_id,
        responsible_executor_id=responsible_executor_id,
        issued_by=issued_by,
        valid_from=valid_from,
        valid_to=valid_to,
        status=WorkOrderStatus.ACTIVE,
        materials=materials,
        tools=tools,
        equipment=equipment,
        special_machinery=special_machinery,
        technological_card_ref=technological_card_ref,
        safety_systems=safety_systems,
        special_conditions=special_conditions,
        work_type_id=work_type_id,
        titul_id=titul_id,
        # Размер бригады на выпуске — база для правила 782н о половине. Если наряд
        # создаётся пустым черновиком, оставляем None; зафиксируется при первом
        # заполнении состава (см. роут /members в webforms.py).
        initial_member_count=(len(member_employee_ids) or None),
    )
    session.add(order)
    session.flush()  # получить order.id до commit

    for employee_id in member_employee_ids:
        session.add(WorkOrderMember(work_order_id=order.id, employee_id=employee_id))

    # Заготовка строк "Ежедневный допуск к работе" на каждый день периода действия —
    # по образцу реального бланка (там наряд на 9 дней = 9 отдельных строк допуска).
    # Пустые (briefing_time/completion_time=NULL) — заполняются по факту каждый день,
    # см. record_daily_admission/record_daily_completion ниже.
    day = valid_from
    while day <= valid_to:
        session.add(WorkOrderDailyAdmission(work_order_id=order.id, admission_date=day))
        day += timedelta(days=1)

    session.commit()
    return order


# ============================================================================
# Пункт 7 наряда-допуска: изменения состава бригады (782н)
# ============================================================================
# Правило 782н: суммарное изменение состава МЕНЕЕ чем на половину первоначального
# разрешено ответственному руководителю по согласованию. Изменение БОЛЕЕ чем
# наполовину (строго больше половины), смена ответственных или условий работы →
# наряд аннулируется, нужен новый. Здесь считаем только члены бригады; смена
# ответственных в наряде не предусмотрена (это всегда новый наряд).

def work_order_change_stats(session: Session, work_order_id: str) -> dict:
    """Статистика изменений состава по наряду для правила половины.
    Возвращает: initial (первоначальный размер), changes (суммарно вводы+выводы),
    limit (макс. допустимое число изменений), exceeded (превышен ли предел),
    at_limit (достигнут ли предел — следующее изменение превысит)."""
    order = session.get(WorkOrder, work_order_id)
    if order is None:
        return {"initial": 0, "changes": 0, "limit": 0, "exceeded": False, "at_limit": True}
    initial = order.initial_member_count
    if not initial:
        # Старый наряд без зафиксированного размера — fallback на текущий состав.
        initial = (
            session.query(WorkOrderMember)
            .filter_by(work_order_id=work_order_id)
            .count()
        )
    changes = (
        session.query(WorkOrderMemberChange)
        .filter_by(work_order_id=work_order_id)
        .count()
    )
    # Разрешено, пока суммарное изменение НЕ превышает половину: changes*2 <= initial.
    # limit — наибольшее число изменений, при котором ещё не «более чем наполовину».
    limit = initial // 2
    exceeded = (changes * 2) > initial
    # at_limit: ещё одно изменение сделает (changes+1)*2 > initial.
    at_limit = ((changes + 1) * 2) > initial
    return {
        "initial": initial,
        "changes": changes,
        "limit": limit,
        "exceeded": exceeded,
        "at_limit": at_limit,
    }


def add_member_change(
    session: Session,
    work_order_id: str,
    employee_id: str,
    change_type: MemberChangeType,
    ordered_by: str,
    created_by: str | None = None,
) -> tuple[bool, str]:
    """Оформляет изменение состава (ввод/вывод) по пункту 7 с проверкой правила
    половины. Возвращает (успех, сообщение). При превышении половины НЕ вносит
    изменение и возвращает (False, требование нового наряда).

    Помимо записи в журнал изменений, фактически меняет состав:
    - added:   добавляет WorkOrderMember, если его ещё нет;
    - removed: удаляет WorkOrderMember этого сотрудника."""
    order = session.get(WorkOrder, work_order_id)
    if order is None:
        return False, "Наряд не найден."

    stats = work_order_change_stats(session, work_order_id)
    # Это изменение — ещё +1 к суммарному. Проверяем, не превысит ли половину.
    if ((stats["changes"] + 1) * 2) > stats["initial"]:
        return False, (
            f"Изменение состава превысит половину первоначального "
            f"({stats['initial']} чел.). По Правилам 782н наряд-допуск в этом "
            f"случае аннулируется — оформите НОВЫЙ наряд-допуск."
        )

    # Фактическое изменение состава.
    existing = (
        session.query(WorkOrderMember)
        .filter_by(work_order_id=work_order_id, employee_id=employee_id)
        .first()
    )
    if change_type == MemberChangeType.ADDED:
        if existing is None:
            session.add(WorkOrderMember(work_order_id=work_order_id, employee_id=employee_id))
    else:  # REMOVED
        if existing is not None:
            session.delete(existing)

    session.add(WorkOrderMemberChange(
        work_order_id=work_order_id,
        employee_id=employee_id,
        change_type=change_type,
        ordered_by=ordered_by,
        created_by=created_by,
    ))
    session.commit()
    action = "введён в состав" if change_type == MemberChangeType.ADDED else "выведен из состава"
    return True, f"Работник {action}. Изменение зафиксировано в пункте 7 наряда."


def get_member_changes(session: Session, work_order_id: str) -> list[WorkOrderMemberChange]:
    """Записи изменений состава по наряду (для пункта 7 бланка и истории в UI)."""
    return (
        session.query(WorkOrderMemberChange)
        .filter_by(work_order_id=work_order_id)
        .order_by(WorkOrderMemberChange.changed_at)
        .all()
    )


# ============================================================================
# ЖУРНАЛ УЧЁТА РАБОТ ПО НАРЯДУ-ДОПУСКУ (Приложение №5 к 782н)
# ============================================================================
# Данные уже есть в WorkOrder/WorkOrderMember — журнал это «вид» существующих
# нарядов с допечаткой партиями и сквозной нумерацией (как у инструктажей).

def get_unprinted_work_orders(session: Session) -> list[WorkOrder]:
    """Наряды, ещё не попавшие в распечатанную партию журнала (journal_row_number пуст)."""
    return (
        session.query(WorkOrder)
        .filter(WorkOrder.is_deleted.is_(False))
        .filter(WorkOrder.journal_row_number.is_(None))
        .order_by(WorkOrder.created_at)
        .all()
    )


def get_last_wo_journal_row_number(session: Session) -> int:
    last = (
        session.query(WorkOrder)
        .filter(WorkOrder.journal_row_number.isnot(None))
        .order_by(WorkOrder.journal_row_number.desc())
        .first()
    )
    return last.journal_row_number if last else 0


def get_journaled_work_orders(session: Session) -> list[WorkOrder]:
    """Наряды, уже внесённые в журнал (есть journal_row_number) — для просмотра
    содержимого журнала на странице, по порядку строк."""
    return (
        session.query(WorkOrder)
        .filter(WorkOrder.journal_row_number.isnot(None))
        .order_by(WorkOrder.journal_row_number)
        .all()
    )


def print_new_wo_journal_entries(session: Session) -> list[WorkOrder]:
    """Допечатать новые наряды в журнал: присваивает сквозные номера строк,
    помечает journal_printed_at. Уже напечатанные наряды не трогает."""
    unprinted = get_unprinted_work_orders(session)
    if not unprinted:
        return []
    next_num = get_last_wo_journal_row_number(session) + 1
    now = datetime.utcnow()
    for wo in unprinted:
        wo.journal_row_number = next_num
        wo.journal_printed_at = now
        next_num += 1
    session.commit()
    return unprinted


# ============================================================================
# ОБЩИЙ ЖУРНАЛ РАБОТ (ОЖР) — внутренний, электронный, заполняется по наряду
# ============================================================================

# Координаты стройплощадки (село Белокаменка, западный берег Кольского залива).
# Погода одинакова в радиусе нескольких км, точности села достаточно для журнала.
_BELOKAMENKA_LAT = 69.08
_BELOKAMENKA_LON = 33.17


def fetch_weather_belokamenka(entry_date: date) -> str | None:
    """Погода на дату по координатам Белокаменки через Open-Meteo (без ключа).
    Формат: температура цифрами + осадки словами, напр. «-8°C (мин -12), снег».
    Минимальная температура показывается всегда — от неё зависит режим зимнего
    бетонирования (СП 70.13330: мин. суточная < 0°C → особый уход).

    Любая ошибка (нет сети, сервис недоступен, дата вне архива) → None: погода
    это удобство, не блокер — запись журнала создаётся и без неё."""
    import json
    import urllib.request
    import urllib.parse

    today = datetime.utcnow().date()
    # Прошлые даты — архивный API; сегодня/будущее — прогнозный.
    if entry_date < today:
        base = "https://archive-api.open-meteo.com/v1/archive"
    else:
        base = "https://api.open-meteo.com/v1/forecast"
    params = {
        "latitude": _BELOKAMENKA_LAT,
        "longitude": _BELOKAMENKA_LON,
        "start_date": entry_date.isoformat(),
        "end_date": entry_date.isoformat(),
        "daily": "temperature_2m_mean,temperature_2m_min,temperature_2m_max,precipitation_sum,snowfall_sum",
        "timezone": "Europe/Moscow",
    }
    url = base + "?" + urllib.parse.urlencode(params)
    try:
        with urllib.request.urlopen(url, timeout=8) as r:
            data = json.load(r)
        daily = data.get("daily", {})
        def _first(key):
            v = daily.get(key)
            return v[0] if isinstance(v, list) and v else None
        t_mean = _first("temperature_2m_mean")
        t_min = _first("temperature_2m_min")
        rain = _first("precipitation_sum")
        snow = _first("snowfall_sum")
        if t_mean is None:
            return None
        if snow and snow > 0:
            precip = "снег"
        elif rain and rain > 0:
            precip = "дождь"
        else:
            precip = "без осадков"
        s = f"{round(t_mean):+d}°C"
        if t_min is not None:
            s += f" (мин {round(t_min):+d})"
        s += f", {precip}"
        return s
    except Exception:
        return None


def create_work_log_entry(
    session: Session, work_order_id: str, entry_date: date, work_done: str,
    weather: str | None = None, note: str | None = None, created_by: str | None = None,
    done_operations: str | None = None,
) -> WorkLogEntry:
    """Создаёт запись ОЖР по наряду. Место/работа/состав НЕ дублируются в запись —
    они берутся из связанного наряда при отображении/печати. Статус — черновик
    (подпись УКЭП ставится отдельно, когда подключён КриптоПро). done_operations —
    JSON-список отмеченных операций вида работ (для прогресса и сборки текста)."""
    entry = WorkLogEntry(
        work_order_id=work_order_id,
        entry_date=entry_date,
        work_done=work_done,
        weather=weather,
        note=note,
        created_by=created_by,
        done_operations=done_operations,
        sign_status=WorkLogSignStatus.DRAFT,
    )
    session.add(entry)
    session.commit()
    return entry


# --- Этапность ОЖР: операции вида работ, прогресс, проверка последовательности ----

def parse_operations(content: str | None) -> list[str]:
    """Операции вида работ из WorkType.content — разбивка по «;». Порядок сохраняется:
    в справочнике операции записаны в технологической последовательности (приёмка →
    заготовка → сборка → установка), это и есть микроэтапность внутри вида."""
    if not content:
        return []
    return [op.strip() for op in content.split(";") if op.strip()]


def get_done_operations_for_order(session: Session, work_order_id: str) -> set[str]:
    """Операции, уже отмеченными в ПРОШЛЫХ записях ОЖР этого наряда (из done_operations).
    По ним экран заполнения понимает прогресс и предлагает следующие невыполненные."""
    done: set[str] = set()
    for e in session.query(WorkLogEntry).filter_by(work_order_id=work_order_id).all():
        if e.done_operations:
            try:
                done.update(json.loads(e.done_operations))
            except Exception:
                pass
    return done


def stage1_documented_for_object(session: Session, titul_id: str | None) -> bool:
    """Критерий (в): по объекту (titul_id) есть хоть одна запись ОЖР этапа 1 (stage_order==1 —
    опалубка/армирование). titul_id пустой → True (объект вписан вручную, межнарядную проверку
    не делаем — не мешаем). Используется для мягкого предупреждения СП 435 перед бетоном."""
    if not titul_id:
        return True
    q = (
        session.query(WorkLogEntry)
        .join(WorkOrder, WorkLogEntry.work_order_id == WorkOrder.id)
        .join(WorkType, WorkOrder.work_type_id == WorkType.id)
        .filter(WorkOrder.titul_id == titul_id, WorkType.stage_order == 1)
    )
    return q.first() is not None


def build_work_done_text(work_type, done_ops: list[str], tools: str | None = None) -> str:
    """Собирает «выполнено за день» из названия вида работ, отмеченных операций и инструментов.
    Инструмент берётся из переданного значения или из WorkType.tools."""
    parts = []
    if work_type is not None and getattr(work_type, "name", None):
        parts.append(work_type.name + ":")
    if done_ops:
        parts.append("; ".join(done_ops))
    text = " ".join(parts).strip() if parts else "; ".join(done_ops or [])
    t = tools if tools is not None else (getattr(work_type, "tools", None) if work_type else None)
    if t and t.strip():
        text = (text + f"\nИнструмент: {t.strip()}").strip()
    return text


def get_work_log_entries(session: Session, work_order_id: str | None = None) -> list[WorkLogEntry]:
    """Записи ОЖР: все или по конкретному наряду, в хронологии по дате записи."""
    q = session.query(WorkLogEntry)
    if work_order_id:
        q = q.filter_by(work_order_id=work_order_id)
    return q.order_by(WorkLogEntry.entry_date, WorkLogEntry.created_at).all()


def delete_work_log_entry(session: Session, entry_id: str) -> bool:
    """Удаление записи ОЖР. Подписанные записи удалять нельзя — только черновики."""
    entry = session.get(WorkLogEntry, entry_id)
    if entry is None:
        return False
    if entry.sign_status == WorkLogSignStatus.SIGNED:
        return False  # подписанное УКЭП не удаляем
    session.delete(entry)
    session.commit()
    return True


def get_unprinted_work_log(session: Session) -> list[WorkLogEntry]:
    return (
        session.query(WorkLogEntry)
        .filter(WorkLogEntry.journal_row_number.is_(None))
        .order_by(WorkLogEntry.entry_date, WorkLogEntry.created_at)
        .all()
    )


def get_last_worklog_row_number(session: Session) -> int:
    last = (
        session.query(WorkLogEntry)
        .filter(WorkLogEntry.journal_row_number.isnot(None))
        .order_by(WorkLogEntry.journal_row_number.desc())
        .first()
    )
    return last.journal_row_number if last else 0


def print_new_work_log_entries(session: Session) -> list[WorkLogEntry]:
    """Допечатать новые записи ОЖР партией со сквозной нумерацией."""
    unprinted = get_unprinted_work_log(session)
    if not unprinted:
        return []
    next_num = get_last_worklog_row_number(session) + 1
    now = datetime.utcnow()
    for e in unprinted:
        e.journal_row_number = next_num
        e.printed_at = now
        next_num += 1
    session.commit()
    return unprinted


def sign_work_log_entry(
    session: Session, entry_id: str, signed_by: str, cert_serial: str, content_hash: str,
) -> bool:
    """Проставляет подпись УКЭП на запись ОЖР. Вызывается ПОСЛЕ того, как клиент
    (КриптоПро browser plug-in) реально подписал содержимое — сюда приходят готовые
    ФИО подписавшего, серийник сертификата и хеш подписанного содержимого. Сервер
    подпись не создаёт, только фиксирует. Заглушка потока — реальный обработчик
    подключится, когда настроим плагин; функция уже готова принять результат."""
    entry = session.get(WorkLogEntry, entry_id)
    if entry is None:
        return False
    entry.sign_status = WorkLogSignStatus.SIGNED
    entry.signed_by = signed_by
    entry.signed_at = datetime.utcnow()
    entry.sign_cert_serial = cert_serial
    entry.content_hash = content_hash
    session.commit()
    return True


def update_work_log_entry(
    session: Session, entry_id: str,
    entry_date: date | None = None, work_done: str | None = None,
    weather: str | None = None, note: str | None = None,
) -> bool:
    """ТЕСТОВОЕ редактирование черновика ОЖР (для отладочного цикла — крутить один документ,
    а не плодить новые). Меняет только записи в статусе DRAFT: подписанное не трогаем.
    На ПРОДЕ эту возможность закрываем — черновик правится до подписи, подпись фиксирует
    содержимое окончательно. entry_date/work_done меняются, если переданы; weather/note
    перезаписываются переданными значениями (в т.ч. на None = очистить)."""
    entry = session.get(WorkLogEntry, entry_id)
    if entry is None or entry.sign_status == WorkLogSignStatus.SIGNED:
        return False
    if entry_date is not None:
        entry.entry_date = entry_date
    if work_done is not None:
        entry.work_done = work_done
    entry.weather = weather
    entry.note = note
    session.commit()
    return True


def unsign_work_log_entry(session: Session, entry_id: str) -> bool:
    """ТЕСТОВАЯ отмена подписи: SIGNED → DRAFT, обнуляет реквизиты подписи (ФИО, серийник,
    хеш, дату). Нужна ТОЛЬКО для отладочного цикла, чтобы гонять один документ туда-сюда.
    На ПРОДЕ убрать вместе с кнопкой — подпись УКЭП необратима."""
    entry = session.get(WorkLogEntry, entry_id)
    if entry is None:
        return False
    entry.sign_status = WorkLogSignStatus.DRAFT
    entry.signed_by = None
    entry.signed_at = None
    entry.sign_cert_serial = None
    entry.content_hash = None
    session.commit()
    return True


def sign_work_order_member(session: Session, work_order_id: str, employee_id: str) -> bool:
    """Подтверждение ознакомления членом бригады (подпись)."""
    member = (
        session.query(WorkOrderMember)
        .filter_by(work_order_id=work_order_id, employee_id=employee_id)
        .first()
    )
    if member is None:
        return False
    member.signed_at = datetime.utcnow()
    session.add(member)
    session.commit()
    return True


def get_daily_admissions(session: Session, work_order_id: str) -> list[WorkOrderDailyAdmission]:
    return (
        session.query(WorkOrderDailyAdmission)
        .filter_by(work_order_id=work_order_id)
        .order_by(WorkOrderDailyAdmission.admission_date)
        .all()
    )


def record_daily_briefing(session: Session, work_order_id: str, admission_date: date,
                           confirmed_by: str) -> bool:
    """Целевой инструктаж на конкретный день выдан — фиксирует время и кто подтвердил
    (ответственный руководитель работ). Одна строка на день, см. WorkOrderDailyAdmission."""
    row = (
        session.query(WorkOrderDailyAdmission)
        .filter_by(work_order_id=work_order_id, admission_date=admission_date)
        .first()
    )
    if row is None:
        return False
    row.briefing_time = datetime.utcnow()
    row.briefing_confirmed_by = confirmed_by
    session.add(row)
    session.commit()
    return True


def record_daily_completion(session: Session, work_order_id: str, admission_date: date,
                             confirmed_by: str) -> bool:
    """Работа за конкретный день закончена, место убрано — фиксирует время и кто
    подтвердил (ответственный исполнитель работ)."""
    row = (
        session.query(WorkOrderDailyAdmission)
        .filter_by(work_order_id=work_order_id, admission_date=admission_date)
        .first()
    )
    if row is None:
        return False
    row.completion_time = datetime.utcnow()
    row.completion_confirmed_by = confirmed_by
    session.add(row)
    session.commit()
    return True


def get_active_work_orders(session: Session) -> list[WorkOrder]:
    today = date.today()
    return (
        session.query(WorkOrder)
        .filter(WorkOrder.is_deleted.is_(False))
        .filter(WorkOrder.status == WorkOrderStatus.ACTIVE)
        .filter(WorkOrder.valid_to >= today)
        .order_by(WorkOrder.valid_from)
        .all()
    )


def get_past_work_orders(session: Session) -> list[WorkOrder]:
    """Архив — наряды с истёкшим сроком (valid_to < сегодня), не удалённые."""
    today = date.today()
    return (
        session.query(WorkOrder)
        .filter(WorkOrder.is_deleted.is_(False))
        .filter(WorkOrder.valid_to < today)
        .order_by(WorkOrder.valid_from.desc(), WorkOrder.number)
        .all()
    )


def get_deleted_work_orders(session: Session) -> list[WorkOrder]:
    """Корзина — мягко удалённые наряды, доступные для восстановления."""
    return (
        session.query(WorkOrder)
        .filter(WorkOrder.is_deleted.is_(True))
        .order_by(WorkOrder.deleted_at.desc())
        .all()
    )


def soft_delete_work_order(session: Session, work_order_id: str, deleted_by: str) -> bool:
    """Мягкое удаление: помечаем is_deleted, наряд уходит в корзину (не стирается)."""
    order = session.get(WorkOrder, work_order_id)
    if order is None or order.is_deleted:
        return False
    order.is_deleted = True
    order.deleted_at = datetime.utcnow()
    order.deleted_by = deleted_by
    session.commit()
    return True


def restore_work_order(session: Session, work_order_id: str) -> bool:
    """Восстановление из корзины: снимаем метку удаления."""
    order = session.get(WorkOrder, work_order_id)
    if order is None or not order.is_deleted:
        return False
    order.is_deleted = False
    order.deleted_at = None
    order.deleted_by = None
    session.commit()
    return True


def hard_delete_work_order(session: Session, work_order_id: str) -> bool:
    """Физическое удаление (стирание из корзины) — только для админа. Каскад снимает
    состав, ежедневные допуски и изменения состава. Необратимо."""
    order = session.get(WorkOrder, work_order_id)
    if order is None:
        return False
    session.delete(order)
    session.commit()
    return True


def close_work_order(session: Session, work_order_id: str) -> bool:
    order = session.get(WorkOrder, work_order_id)
    if order is None:
        return False
    order.status = WorkOrderStatus.CLOSED
    session.add(order)
    session.commit()
    return True


# ================= Инструктажи =================

INSTRUCTION_LABELS = {
    InstructionType.INTRODUCTORY: "Вводный",
    InstructionType.PRIMARY_WORKPLACE: "Первичный на рабочем месте",
    InstructionType.REPEATED: "Повторный",
    InstructionType.UNSCHEDULED: "Внеплановый",
    InstructionType.TARGETED: "Целевой",
}

# Строк на страницу при допечатке журнала — [Предполагаю] не согласовано явно,
# разумный дефолт под таблицу А4 с таким набором граф. Легко поменять.
JOURNAL_ROWS_PER_PAGE = 20

# Порог "критично" для непроведённых вводного/первичного инструктажей: сколько
# дней после ДАТЫ НАЧАЛА РАБОТЫ считается некритичной просрочкой, после чего —
# критичной. Согласовано явно: 5 дней. Стадии: "overdue" (дата начала прошла,
# 0..5 дней) → "critical" (> 5 дней). Стадия "заранее" для этих двух видов на
# практике не возникает, пока даты начала в прошлом (см. диалог) — включится
# сама, когда появятся будущие даты начала работы.
INSTRUCTION_OVERDUE_CRITICAL_DAYS = 5

# Виды инструктажа, проведение которых система ТРЕБУЕТ по дате начала работы
# (разовые, привязаны к приёму). Повторный/внеплановый/целевой сюда НЕ входят:
# повторный периодический (отложен), внеплановый/целевой — событийные, планового
# срока по календарю у них нет.
REQUIRED_INSTRUCTION_TYPES = (
    InstructionType.INTRODUCTORY,
    InstructionType.PRIMARY_WORKPLACE,
)


def get_employees_needing_instruction(
    session: Session, instruction_type: InstructionType
) -> list[Employee]:
    """Активные сотрудники, у кого известна дата начала работы (дата договора,
    а если её ещё нет — дата въезда), но инструктажа ДАННОГО типа ещё нет ни
    одного. Обобщение прежней get_employees_needing_introductory на любой тип —
    чтобы вводный и первичный проверялись одной логикой, а не двумя копиями."""
    existing_ids = {
        i.employee_id for i in
        session.query(Instruction).filter_by(type=instruction_type).all()
    }
    employees = (
        session.query(Employee)
        .filter(Employee.contract_end_date.is_(None))
        .filter((Employee.contract_date.isnot(None)) | (Employee.entry_date.isnot(None)))
        .all()
    )
    return [e for e in employees if e.id not in existing_ids]


def get_employees_needing_introductory(session: Session) -> list[Employee]:
    """Обёртка над get_employees_needing_instruction для вводного — сохранена,
    чтобы не менять вызовы в webforms.py (кнопка «Заполнить вводный всем»)."""
    return get_employees_needing_instruction(session, InstructionType.INTRODUCTORY)


def get_instruction_compliance_gaps(session: Session) -> list[dict]:
    """Пробелы по ОБЯЗАТЕЛЬНЫМ инструктажам (вводный + первичный на рабочем месте):
    активный сотрудник, у которого дата начала работы уже наступила, а инструктаж
    этого типа не проведён. Для дашборда веба, раздела «Требует внимания» в боте
    и утренней рассылки «до устранения».

    Стадия:
      "overdue"  — дата начала прошла, 0..INSTRUCTION_OVERDUE_CRITICAL_DAYS дней;
      "critical" — прошло больше порога.
    Сотрудники с ещё НЕ наступившей датой начала (будущий приём) сюда не попадают —
    у них обязанность ещё не возникла (для них позже естественно оживёт стадия
    «заранее», её добавим, когда появятся будущие даты)."""
    today = date.today()
    gaps: list[dict] = []
    for itype in REQUIRED_INSTRUCTION_TYPES:
        for emp in get_employees_needing_instruction(session, itype):
            start_date = emp.contract_date or emp.entry_date
            if start_date is None or start_date > today:
                continue  # дата начала не наступила — обязанности ещё нет
            days_overdue = (today - start_date).days
            stage = "critical" if days_overdue > INSTRUCTION_OVERDUE_CRITICAL_DAYS else "overdue"
            gaps.append({
                "employee_id": emp.id,
                "name": emp.full_name,
                "instruction_type": itype,
                "type_label": INSTRUCTION_LABELS.get(itype, itype.value),
                "start_date": start_date,
                "days_overdue": days_overdue,
                "stage": stage,
            })
    # Критичные первыми, внутри — по возрастанию давности (свежие сверху),
    # чтобы в рассылке/на дашборде взгляд цеплялся за самое горящее.
    gaps.sort(key=lambda g: (g["stage"] != "critical", g["days_overdue"]))
    return gaps


def auto_create_instructions(
    session: Session, instruction_type: InstructionType, conducted_by: str
) -> list[Instruction]:
    """Заводит инструктаж ЗАДАННОГО типа для ВСЕХ сотрудников, у кого его ещё нет —
    разом, датой начала работы каждого (дата договора, если пусто — дата въезда),
    НЕ сегодняшней датой (чтобы порядок строк в журнале совпал с хронологией приёма).

    Используется для вводного и первичного на рабочем месте — оба разовые, оба
    проводятся в день начала работы (по факту вместе, но регистрируются в РАЗНЫХ
    журналах со своей сквозной нумерацией на каждый тип).

    Защита от гонки: коммит ПО ОДНОМУ сотруднику. Если двойное нажатие/параллельный
    запрос создаёт дубль — unique-индекс в БД отклонит именно эту вставку
    (uq_intro_once для вводного, uq_primary_workplace_once для первичного), не
    обрушив пачку остальных. ВАЖНО: для новых типов, требующих уникальности
    "один на сотрудника", такой индекс должен существовать в БД — иначе защиты нет."""
    from sqlalchemy.exc import IntegrityError

    employees = get_employees_needing_instruction(session, instruction_type)
    created = []
    for e in employees:
        start_date = e.contract_date or e.entry_date
        instr = Instruction(
            employee_id=e.id,
            type=instruction_type,
            conducted_by=conducted_by,
            conducted_at=datetime.combine(start_date, datetime.min.time()),
        )
        session.add(instr)
        try:
            session.commit()
            created.append(instr)
        except IntegrityError:
            session.rollback()  # уже есть (гонка/повторное нажатие) — пропускаем, не падаем
    return created


def auto_create_introductory_instructions(session: Session, conducted_by: str) -> list[Instruction]:
    """Обёртка над auto_create_instructions для вводного — сохранена, чтобы не
    менять существующий вызов в webforms.py (кнопка «Заполнить вводный всем»)."""
    return auto_create_instructions(session, InstructionType.INTRODUCTORY, conducted_by)


def get_unprinted_instructions(session: Session, instruction_type: InstructionType) -> list[Instruction]:
    return (
        session.query(Instruction)
        .filter_by(type=instruction_type, printed_at=None)
        .order_by(Instruction.conducted_at)
        .all()
    )


def get_last_journal_row_number(session: Session, instruction_type: InstructionType) -> int:
    last = (
        session.query(Instruction)
        .filter_by(type=instruction_type)
        .filter(Instruction.journal_row_number.isnot(None))
        .order_by(Instruction.journal_row_number.desc())
        .first()
    )
    return last.journal_row_number if last else 0


def print_new_journal_entries(session: Session, instruction_type: InstructionType) -> list[Instruction]:
    """"Допечатать новые записи" — присваивает номера строк последовательно, продолжая
    с последнего уже выданного номера (отдельная нумерация на каждый InstructionType —
    это разные физические журналы), помечает printed_at. Старые (уже напечатанные и
    подшитые) записи не трогает и не перепечатывает — см. договорённость в чате."""
    unprinted = get_unprinted_instructions(session, instruction_type)
    if not unprinted:
        return []
    next_num = get_last_journal_row_number(session, instruction_type) + 1
    now = datetime.utcnow()
    for instr in unprinted:
        instr.journal_row_number = next_num
        instr.printed_at = now
        next_num += 1
    session.commit()
    return unprinted


def get_sheet_count(session: Session, instruction_type: InstructionType) -> int:
    """Сколько партий (= "листов") уже когда-либо напечатано для этого типа —
    приближённо, по числу различных printed_at. [Предполагаю] это не полный
    учёт физических листов бумаги, просто счётчик распечаток."""
    rows = (
        session.query(Instruction.printed_at)
        .filter_by(type=instruction_type)
        .filter(Instruction.printed_at.isnot(None))
        .distinct()
        .all()
    )
    return len(rows)


def create_instruction(
    session: Session, employee_id: str, instruction_type: InstructionType,
    conducted_by: str, topic: str | None = None, next_due_date: date | None = None,
) -> Instruction:
    instr = Instruction(
        employee_id=employee_id,
        type=instruction_type,
        topic=topic,
        conducted_by=conducted_by,
        next_due_date=next_due_date,
    )
    session.add(instr)
    session.commit()
    return instr


def confirm_instruction(session: Session, instruction_id: str) -> bool:
    """Работник подтвердил, что ознакомлен (подпись)."""
    instr = session.get(Instruction, instruction_id)
    if instr is None:
        return False
    instr.employee_confirmed_at = datetime.utcnow()
    session.add(instr)
    session.commit()
    return True


def get_instructions_for_employee(session: Session, employee_id: str) -> list[Instruction]:
    return (
        session.query(Instruction)
        .filter_by(employee_id=employee_id)
        .order_by(Instruction.conducted_at.desc())
        .all()
    )


def _set_cell(cell, text: str, size: int = 8, bold: bool = False) -> None:
    """Текст ячейки с уменьшенным шрифтом — чтобы длинные ФИО помещались в одну
    строку при печати, не переносились на две-три (см. договорённость)."""
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold


def get_journal_started_at(session: Session, instruction_type: InstructionType) -> date | None:
    """Дата самой ранней записи в журнале этого типа — для графы "Начат" на обложке.
    Не дата печати текущей партии, а дата первой КОГДА-ЛИБО занесённой записи."""
    first = (
        session.query(Instruction)
        .filter_by(type=instruction_type)
        .order_by(Instruction.conducted_at)
        .first()
    )
    return first.conducted_at.date() if first else None


def _set_a4(doc) -> None:
    """Явный формат A4 — python-docx по умолчанию создаёт Letter (US), не A4,
    для русских документов это неверно (см. замечание после первой распечатки)."""
    section = doc.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)


def generate_instruction_journal_docx(
    instructions: list[Instruction], instruction_type: InstructionType,
    org_name: str, order_ref: str, journal_number: int = 1,
    started_at: date | None = None, sheet_number: int = 1, total_sheets: int = 1,
    output_dir: str = "/tmp",
) -> str:
    """
    УСТАРЕЛО (2026-07): журнал переведён на Excel — см.
    generate_instruction_journal_xlsx ниже, роут /production/instructions/print
    в webforms.py теперь вызывает xlsx-версию. Эта docx-функция ОСТАВЛЕНА
    НАМЕРЕННО как быстрый откат: если xlsx на реальном принтере ляжет криво,
    достаточно вернуть вызов *_docx в одном роуте, не переписывая ничего с нуля.
    УДАЛИТЬ, когда xlsx-версия подтверждена на живой печати.
    ------------------------------------------------------------------------
    Печать партии журнала инструктажей. 2026-07 (третий заход) — переписано
    под официальный скелет таблицы, который прислал пользователь ("Журнал
    регистрации вводного инструктажа"), а не по собственной реконструкции
    ГОСТ 12.0.004-2015. Отличия от предыдущей версии:
    - Обложка — не текстом, а таблицей: Начат/Окончен/Количество листов/Лист.
    - НЕТ столбца "№" — в присланном скелете его нет вообще (номер строки
      всё равно хранится в БД для сквозной нумерации, просто не печатается
      отдельной графой, раз в образце так).
    - Порядок подписей: СНАЧАЛА инструктирующего, ПОТОМ инструктируемого
      (было наоборот).
    - "Дата рождения" — полная дата, не только год (в Employee есть полная
      дата, раньше зря обрезал до года).
    - Двухуровневая шапка с объединёнными ячейками ("Сведения об
      инструктируемом" на 3 подстолбца, "Подпись" на 2 подстолбца) — как в
      присланном образце.

    sheet_number/total_sheets — для графы "Лист"/"Количество листов" на
    обложке. [Предполагаю] total_sheets = сколько партий уже напечатано
    включая эту — простое приближение, не полноценный учёт физических
    листов бумаги (это была бы отдельная задача подсчёта).
    """
    doc = Document()
    _set_a4(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    label = INSTRUCTION_LABELS.get(instruction_type, instruction_type.value)
    run = title.add_run(f"ЖУРНАЛ РЕГИСТРАЦИИ ИНСТРУКТАЖА ({label.upper()})")
    run.bold = True
    run.font.size = Pt(14)

    doc.add_paragraph(f"Организация: {org_name}          Журнал № {journal_number}")

    # Обложка: Начат | значение | Окончен | значение | Количество листов | значение | Лист | значение
    started_str = started_at.strftime("%d.%m.%Y") if started_at else "—"
    cover = doc.add_table(rows=1, cols=8)
    cover.style = "Table Grid"
    cover_cells = cover.rows[0].cells
    cover_values = [
        ("Начат", started_str), ("Окончен", "—"),
        ("Количество листов", str(total_sheets)), ("Лист", str(sheet_number)),
    ]
    for i, (lbl, val) in enumerate(cover_values):
        _set_cell(cover_cells[i * 2], lbl, size=8, bold=True)
        _set_cell(cover_cells[i * 2 + 1], val, size=8)

    doc.add_paragraph()

    # Основная таблица — двухуровневая шапка с объединением ячеек, 10 столбцов:
    # 0 № сквозной | 1 № на листе | 2 Дата проведения |
    # 3-5 Сведения об инструктируемом (ФИО/дата рожд./профессия) |
    # 6 Подразделение | 7 ФИО+должность инструктирующего | 8-9 Подпись (инстр./инстр-емого)
    table = doc.add_table(rows=2, cols=10)
    table.style = "Table Grid"
    table.autofit = False
    r1 = table.rows[0].cells
    r2 = table.rows[1].cells

    # Вертикальное объединение (2 строки шапки в одну ячейку) для столбцов без подстолбцов.
    for col in (0, 1, 2, 6, 7):
        r1[col].merge(r2[col])
    _set_cell(r1[0], "№ сквозной", size=7, bold=True)
    _set_cell(r1[1], "№ на листе", size=7, bold=True)
    _set_cell(r1[2], "Дата проведения", size=7, bold=True)
    _set_cell(r1[6], "Наименование структурного подразделения, в которое направлен инструктируемый",
              size=7, bold=True)
    _set_cell(r1[7], "Фамилия, имя, отчество, должность инструктирующего", size=7, bold=True)

    # Горизонтальное объединение верхней строки для "Сведения об инструктируемом" (3-5)
    # и "Подпись" (8-9), с подписанными подстолбцами во второй строке.
    r1[3].merge(r1[4]).merge(r1[5])
    _set_cell(r1[3], "Сведения об инструктируемом", size=7, bold=True)
    _set_cell(r2[3], "Фамилия, имя, отчество", size=7, bold=True)
    _set_cell(r2[4], "Дата рождения", size=7, bold=True)
    _set_cell(r2[5], "Профессия, должность", size=7, bold=True)

    r1[8].merge(r1[9])
    _set_cell(r1[8], "Подпись", size=7, bold=True)
    _set_cell(r2[8], "Инструктирующего", size=7, bold=True)
    _set_cell(r2[9], "Инструктируемого", size=7, bold=True)

    # Номер на листе — позиция внутри ТЕКУЩЕЙ партии (1..JOURNAL_ROWS_PER_PAGE),
    # не сквозной. Корректно, только если предыдущая партия допечатана прочерками
    # ровно до конца страницы (см. довесок ниже) — тогда каждая новая партия
    # гарантированно начинается с начала листа, позиция внутри нее = позиция на листе.
    for i, instr in enumerate(instructions):
        row = table.add_row().cells
        emp = instr.employee
        birth_str = emp.birth_date.strftime("%d.%m.%Y") if emp and emp.birth_date else ""
        sheet_pos = (i % JOURNAL_ROWS_PER_PAGE) + 1
        _set_cell(row[0], str(instr.journal_row_number))
        _set_cell(row[1], str(sheet_pos))
        _set_cell(row[2], instr.conducted_at.strftime("%d.%m.%Y"))
        _set_cell(row[3], emp.full_name if emp else "?")
        _set_cell(row[4], birth_str)
        _set_cell(row[5], (emp.position or "") if emp else "")  # из карточки, если заполнено
        _set_cell(row[6], (emp.subdivision or "") if emp else "")  # из карточки, если заполнено
        _set_cell(row[7], instr.conducted_by)
        _set_cell(row[8], "")  # подпись инструктирующего — от руки на распечатке
        _set_cell(row[9], "")  # подпись инструктируемого — от руки на распечатке

    # Довесок прочерками до конца страницы — визуальный, не данные (не создаёт
    # записей Instruction, следующая допечатка продолжит нумерацию с реального
    # следующего номера, не с номера довеска).
    remainder = len(instructions) % JOURNAL_ROWS_PER_PAGE
    if remainder != 0:
        pad_rows = JOURNAL_ROWS_PER_PAGE - remainder
        for _ in range(pad_rows):
            row = table.add_row().cells
            for cell in row:
                _set_cell(cell, "—")

    footer = doc.add_paragraph()
    footer_run = footer.add_run(
        f"Журнал ведётся по рекомендуемой форме ГОСТ 12.0.004-2015. Порядок регистрации "
        f"определён работодателем самостоятельно (п. 88 Правил №2464 от 24.12.2021; "
        f"разъяснение Роструда №15-2/В-1677 от 30.05.2022) — {order_ref}. "
        f"Незаполненные строки погашены прочерком. Подписи — собственноручные. "
        f"Сроки по закону: вводный — в день начала работы; первичный на рабочем месте — до "
        f"допуска к самостоятельной работе; повторный — не реже 1 раза в 6–12 мес.; "
        f"внеплановый/целевой — по факту события."
    )
    footer_run.font.size = Pt(7)
    footer_run.italic = True

    path = f"{output_dir}/journal_{instruction_type.value}_{datetime.utcnow():%Y%m%d%H%M%S}.docx"
    doc.save(path)
    return path


# Границы/шрифт/выравнивание журнала — вынесены в модульные константы, чтобы
# не пересоздавать объекты openpyxl на каждую ячейку (Border/Font неизменяемы,
# один экземпляр можно назначать множеству ячеек).
_XL_THIN = Side(style="thin", color="000000")
_XL_BORDER = Border(left=_XL_THIN, right=_XL_THIN, top=_XL_THIN, bottom=_XL_THIN)
_XL_FONT = Font(name="Times New Roman", size=9)
_XL_FONT_BOLD = Font(name="Times New Roman", size=9, bold=True)
_XL_FONT_SMALL = Font(name="Times New Roman", size=8)
_XL_CENTER = Alignment(horizontal="center", vertical="center", wrap_text=True)
_XL_LEFT = Alignment(horizontal="left", vertical="center", wrap_text=True)


def _xl_cell(ws, row: int, col: int, value, *, bold: bool = False,
             small: bool = False, left: bool = False, border: bool = True):
    """Записать значение в ячейку с единым стилем журнала. Возвращает ячейку,
    чтобы вызывающий код мог при желании доопределить (например, снять границу).

    ВНИМАНИЕ: не вызывать для НЕ-левых-верхних ячеек объединённого диапазона —
    у openpyxl они становятся MergedCell с read-only .value. Для обрамления
    таких ячеек использовать _xl_border_range после merge_cells."""
    cell = ws.cell(row=row, column=col, value=value)
    cell.font = _XL_FONT_BOLD if bold else (_XL_FONT_SMALL if small else _XL_FONT)
    cell.alignment = _XL_LEFT if left else _XL_CENTER
    if border:
        cell.border = _XL_BORDER
    return cell


def _xl_border_range(ws, r1: int, c1: int, r2: int, c2: int) -> None:
    """Проставить границу _XL_BORDER на КАЖДУЮ ячейку прямоугольника r1..r2 × c1..c2,
    включая MergedCell (им нельзя писать .value, но .border — можно). Нужно, чтобы
    объединённые ячейки шапки печатались обрамлёнными по всем внутренним линиям."""
    for r in range(r1, r2 + 1):
        for c in range(c1, c2 + 1):
            ws.cell(row=r, column=c).border = _XL_BORDER


def generate_wo_journal_xlsx(
    orders: list[WorkOrder], org_name: str, output_dir: str = "/tmp",
) -> str:
    """Журнал учёта работ по наряду-допуску (Приложение №5 к 782н), xlsx.
    Графы формы: № наряда, место и наименование работы, ответственный исполнитель
    (с группой), члены бригады (с группами), кто выдал наряд, даты. Печатается
    партией уже пронумерованных нарядов."""
    wb = Workbook()
    ws = wb.active
    ws.title = "Журнал нарядов"
    ws.page_setup.orientation = "landscape"

    _xl_cell(ws, 1, 1, "ЖУРНАЛ УЧЁТА РАБОТ ПО НАРЯДУ-ДОПУСКУ", bold=True, border=False)
    _xl_cell(ws, 2, 1, org_name, small=True, border=False)
    _xl_cell(ws, 3, 1, "Приложение № 5 к Правилам по охране труда при работе на высоте (Приказ Минтруда 782н)",
             small=True, border=False)

    headers = ["№ строки", "№ наряда", "Место и наименование работы",
               "Ответственный исполнитель (ФИО, группа)", "Члены бригады (ФИО, группа)",
               "Наряд выдал (ФИО)", "Начало работ", "Окончание работ"]
    hr = 5
    for i, h in enumerate(headers, 1):
        _xl_cell(ws, hr, i, h, bold=True, small=True)
    widths = [8, 12, 30, 26, 34, 20, 14, 14]
    for i, w in enumerate(widths, 1):
        ws.column_dimensions[chr(64 + i)].width = w

    r = hr + 1
    for wo in orders:
        ex = getattr(wo, "responsible_executor", None)
        ex_grp = getattr(ex, "height_safety_group", "") or ""
        ex_name = (getattr(ex, "full_name", "") or "") + (f", {ex_grp}" if ex_grp else "")
        sup = getattr(wo, "responsible_supervisor", None)
        members = getattr(wo, "members", []) or []
        crew = "; ".join(
            (getattr(getattr(m, "employee", None), "full_name", "") or "")
            + (f" ({getattr(getattr(m, 'employee', None), 'height_safety_group', '') or ''})"
               if getattr(getattr(m, "employee", None), "height_safety_group", "") else "")
            for m in members
        )
        place = f"{wo.location} — {wo.work_description}"
        start = f"{wo.valid_from:%d.%m.%Y}" if wo.valid_from else ""
        end = f"{wo.valid_to:%d.%m.%Y}" if wo.valid_to else ""
        row = [str(wo.journal_row_number or ""), wo.number or "", place, ex_name, crew,
               getattr(sup, "full_name", "") or wo.issued_by or "", start, end]
        for i, val in enumerate(row, 1):
            _xl_cell(ws, r, i, val, small=True, left=(i in (3, 4, 5)))
        r += 1

    _xl_cell(ws, r + 1, 1, f"Внесено записей: {len(orders)}", bold=True, border=False)
    path = f"{output_dir}/journal_naryadov.xlsx"
    wb.save(path)
    return path


def generate_work_log_xlsx(
    entries: list[WorkLogEntry], org_name: str, output_dir: str = "/tmp",
) -> str:
    """Общий журнал работ (ОЖР), xlsx. Каждая строка — день работ по наряду:
    дата, № наряда, место/работа (из наряда), что сделано за день, погода,
    подпись (статус УКЭП). Печатается партией пронумерованных записей."""
    wb = Workbook()
    ws = wb.active
    ws.title = "Общий журнал работ"
    ws.page_setup.orientation = "landscape"

    _xl_cell(ws, 1, 1, "ОБЩИЙ ЖУРНАЛ РАБОТ", bold=True, border=False)
    _xl_cell(ws, 2, 1, org_name, small=True, border=False)

    headers = ["№ строки", "Дата", "№ наряда", "Место и работа (по наряду)",
               "Выполнено за день", "Погодные условия", "Внёс", "Подпись (УКЭП)"]
    hr = 4
    for i, h in enumerate(headers, 1):
        _xl_cell(ws, hr, i, h, bold=True, small=True)
    widths = [8, 12, 12, 30, 40, 18, 18, 22]
    for i, w in enumerate(widths, 1):
        ws.column_dimensions[chr(64 + i)].width = w

    r = hr + 1
    for e in entries:
        wo = getattr(e, "work_order", None)
        place = f"{getattr(wo, 'location', '')} — {getattr(wo, 'work_description', '')}" if wo else ""
        wo_num = getattr(wo, "number", "") if wo else ""
        if e.sign_status == WorkLogSignStatus.SIGNED:
            sign = f"Подписано: {e.signed_by or ''} {e.signed_at:%d.%m.%Y}" if e.signed_at else "Подписано"
        else:
            sign = "черновик (не подписано)"
        row = [str(e.journal_row_number or ""), f"{e.entry_date:%d.%m.%Y}", wo_num, place,
               e.work_done or "", e.weather or "", e.created_by or "", sign]
        for i, val in enumerate(row, 1):
            _xl_cell(ws, r, i, val, small=True, left=(i in (4, 5)))
        r += 1

    _xl_cell(ws, r + 1, 1, f"Внесено записей: {len(entries)}", bold=True, border=False)
    path = f"{output_dir}/obshchiy_zhurnal_rabot.xlsx"
    wb.save(path)
    return path


def generate_instruction_journal_xlsx(
    instructions: list[Instruction], instruction_type: InstructionType,
    org_name: str, order_ref: str, journal_number: int = 1,
    started_at: date | None = None, sheet_number: int = 1, total_sheets: int = 1,
    output_dir: str = "/tmp",
) -> str:
    """
    Excel-версия печати журнала инструктажей. Тот же официальный скелет таблицы
    (10 столбцов, двухуровневая шапка, обложка Начат/Окончен), печатается для
    любого типа (вводный, первичный на рабочем месте) — заголовок берётся из
    INSTRUCTION_LABELS по instruction_type, отдельный шаблон под каждый тип не нужен.

    2026-07 (переработка разбивки): шапка и разбивка на листы делаются ФИЗИЧЕСКИ,
    а не через print_title_rows. Причина — print_title_rows это инструкция для
    драйвера печати, её НЕ рендерят мобильные просмотрщики (в т.ч. просмотр вложения
    в MAX): при работе с телефона пользователь видел один сплошной лист без повтора
    шапки. Теперь блок шапки вставляется реальными строками заново на каждом листе
    + жёсткий разрыв страницы (row_breaks) — шапка видна и в просмотре, и в печати,
    а число строк на лист фиксировано (не зависит от драйвера и от того, есть ли
    сверху блок обложки).

    Разбивка: обложка (название/организация/Начат-Окончен) — ТОЛЬКО на первом листе,
    поэтому на нём строк данных меньше (ROWS_FIRST_PAGE), на последующих больше
    (ROWS_OTHER_PAGES) — так на альбомном A4 каждый лист заполнен без переполнения.
    Двухуровневая шапка таблицы повторяется на КАЖДОМ листе.

    Под таблицей — итог "Внесено записей: N (строки M–K)": N в этом файле (партия
    печати), M–K — диапазон сквозных номеров. Между блоком заголовка и таблицей —
    пустая строка-разделитель, чтобы обложка визуально не сливалась с данными.
    """
    # Сколько строк ДАННЫХ помещается на альбомном A4:
    # первый лист несёт блок обложки сверху (название+организация+Начат/Окончен+
    # разделитель), поэтому данных на нём меньше; последующие — только шапка таблицы.
    ROWS_FIRST_PAGE = 17
    ROWS_OTHER_PAGES = 21

    wb = Workbook()
    ws = wb.active
    ws.title = "Журнал"

    # --- Печать: альбомная A4, впис по ширине ---
    ws.page_setup.orientation = "landscape"
    ws.page_setup.paperSize = 9  # A4
    ws.page_setup.fitToWidth = 1
    ws.page_setup.fitToHeight = 0  # по высоте — сколько листов нужно, не сжимать
    ws.sheet_properties.pageSetUpPr.fitToPage = True

    NCOLS = 10  # столбцы A..J
    label = INSTRUCTION_LABELS.get(instruction_type, instruction_type.value)

    page_break_rows: list[int] = []

    def _write_cover(r: int, page_no: int) -> int:
        """Блок обложки (ТОЛЬКО лист 1): название журнала, организация, строка
        Начат/Окончен/Количество листов/Лист, затем пустая строка-разделитель.
        "Количество листов" — пусто (общее число листов журнала неизвестно, пока
        журнал открыт; заполняется от руки при сшивке). "Лист" — сквозной номер
        листа (на листе 1 это 1). На листах 2+ Начат/Окончен НЕ повторяются —
        там только компактная строка "Лист N" (см. _write_sheet_label)."""
        ws.merge_cells(start_row=r, start_column=1, end_row=r, end_column=NCOLS)
        c = ws.cell(row=r, column=1, value=f"ЖУРНАЛ РЕГИСТРАЦИИ ИНСТРУКТАЖА ({label.upper()})")
        c.font = Font(name="Times New Roman", size=14, bold=True)
        c.alignment = Alignment(horizontal="center", vertical="center")
        ws.merge_cells(start_row=r + 1, start_column=1, end_row=r + 1, end_column=NCOLS)
        c = ws.cell(row=r + 1, column=1,
                    value=f"Организация: {org_name}          Журнал № {journal_number}")
        c.font = _XL_FONT
        c.alignment = Alignment(horizontal="left", vertical="center")
        started_str = started_at.strftime("%d.%m.%Y") if started_at else "—"
        cover = [
            ("Начат", started_str), ("Окончен", "—"),
            ("Количество листов", ""), ("Лист", str(page_no)),
        ]
        col = 1
        for lbl, val in cover:
            _xl_cell(ws, r + 2, col, lbl, bold=True, small=True)
            _xl_cell(ws, r + 2, col + 1, val, small=True)
            col += 2
        _xl_cell(ws, r + 2, 9, "", small=True)
        _xl_cell(ws, r + 2, 10, "", small=True)
        return r + 4  # r, r+1, r+2 заняты + r+3 пустой разделитель → следующая r+4

    def _write_sheet_label(r: int, page_no: int) -> int:
        """Компактная строка "Лист N" для листов 2+ (без Начат/Окончен — они только
        на листе 1). Справа над таблицей. Возвращает следующую строку."""
        ws.merge_cells(start_row=r, start_column=1, end_row=r, end_column=NCOLS)
        c = ws.cell(row=r, column=1, value=f"Лист {page_no}")
        c.font = _XL_FONT_BOLD
        c.alignment = Alignment(horizontal="right", vertical="center")
        return r + 1

    def _write_table_header(r: int) -> int:
        """Двухуровневая шапка таблицы (2 строки). Вставляется на КАЖДОМ листе.
        Возвращает первую строку данных (r + 2)."""
        h1, h2 = r, r + 1
        single = {
            1: "№ сквозной",
            2: "№ на листе",
            3: "Дата проведения",
            7: "Наименование структурного подразделения, в которое направлен инструктируемый",
            8: "Фамилия, имя, отчество, должность инструктирующего",
        }
        for col_idx, text in single.items():
            ws.merge_cells(start_row=h1, start_column=col_idx, end_row=h2, end_column=col_idx)
            _xl_cell(ws, h1, col_idx, text, bold=True, small=True)
            _xl_border_range(ws, h1, col_idx, h2, col_idx)
        ws.merge_cells(start_row=h1, start_column=4, end_row=h1, end_column=6)
        _xl_cell(ws, h1, 4, "Сведения об инструктируемом", bold=True, small=True)
        _xl_border_range(ws, h1, 4, h1, 6)
        _xl_cell(ws, h2, 4, "Фамилия, имя, отчество", bold=True, small=True)
        _xl_cell(ws, h2, 5, "Дата рождения", bold=True, small=True)
        _xl_cell(ws, h2, 6, "Профессия, должность", bold=True, small=True)
        ws.merge_cells(start_row=h1, start_column=9, end_row=h1, end_column=10)
        _xl_cell(ws, h1, 9, "Подпись", bold=True, small=True)
        _xl_border_range(ws, h1, 9, h1, 10)
        _xl_cell(ws, h2, 9, "Инструктирующего", bold=True, small=True)
        _xl_cell(ws, h2, 10, "Инструктируемого", bold=True, small=True)
        ws.row_dimensions[h1].height = 30
        ws.row_dimensions[h2].height = 42
        return r + 2

    def _write_data_row(r: int, pos_on_page: int, instr) -> None:
        emp = instr.employee
        birth_str = emp.birth_date.strftime("%d.%m.%Y") if emp and emp.birth_date else ""
        _xl_cell(ws, r, 1, str(instr.journal_row_number))
        _xl_cell(ws, r, 2, str(pos_on_page + 1))  # № на листе (позиция на текущем листе)
        _xl_cell(ws, r, 3, instr.conducted_at.strftime("%d.%m.%Y"))
        _xl_cell(ws, r, 4, emp.full_name if emp else "?", left=True)
        _xl_cell(ws, r, 5, birth_str)
        _xl_cell(ws, r, 6, (emp.position or "") if emp else "", left=True)
        _xl_cell(ws, r, 7, (emp.subdivision or "") if emp else "", left=True)
        _xl_cell(ws, r, 8, instr.conducted_by, left=True)
        _xl_cell(ws, r, 9, "")
        _xl_cell(ws, r, 10, "")
        ws.row_dimensions[r].height = 22

    def _write_page_summary(r: int, cumulative: int, last_seq) -> None:
        """Итоговая строка В КОНЦЕ ЛИСТА (смысл 1): накопительный счётчик записей
        ТЕКУЩЕЙ ПАРТИИ на конец этого листа + реальный последний сквозной номер из
        БД. При допечатке партиями эти числа расходятся (партия из 8 записей может
        нести сквозные 48–55), поэтому пишем оба — "в этой партии" отвечает на
        "сколько внесла распечатка", сквозной — на "с какого номера продолжать"."""
        seq_part = f" · по строку {last_seq}" if last_seq is not None else ""
        text = f"Внесено в этой партии: {cumulative}{seq_part}"
        ws.merge_cells(start_row=r, start_column=1, end_row=r, end_column=NCOLS)
        c = ws.cell(row=r, column=1, value=text)
        c.font = _XL_FONT_BOLD
        c.alignment = Alignment(horizontal="left", vertical="center")

    # --- Раскладка по листам ---
    row = 1
    idx = 0
    n = len(instructions)
    page_num = 0
    while idx < n or page_num == 0:
        page_num += 1
        if page_num == 1:
            row = _write_cover(row, page_num)  # обложка + "Лист 1" в её ячейке
            capacity = ROWS_FIRST_PAGE
        else:
            row = _write_sheet_label(row, page_num)  # компактная "Лист N", без обложки
            capacity = ROWS_OTHER_PAGES
        row = _write_table_header(row)
        pos_on_page = 0
        last_seq_on_page = None
        while pos_on_page < capacity and idx < n:
            _write_data_row(row, pos_on_page, instructions[idx])
            last_seq_on_page = instructions[idx].journal_row_number
            row += 1
            idx += 1
            pos_on_page += 1
        # Итог в конце листа: накопительно по партии (idx = сколько уже разложено)
        # + последний сквозной номер этого листа.
        if n > 0:
            _write_page_summary(row, idx, last_seq_on_page)
            row += 1
        # Разрыв страницы после строки итога, если впереди ещё есть записи.
        if idx < n:
            page_break_rows.append(row - 1)
        if n == 0:
            break

    for br in page_break_rows:
        ws.row_breaks.append(Break(id=br))

    # --- Общий футер (только под последним листом) ---
    row += 1  # пустая строка перед футером

    footer_text = (
        f"Журнал ведётся по рекомендуемой форме ГОСТ 12.0.004-2015. Порядок регистрации "
        f"определён работодателем самостоятельно (п. 88 Правил №2464 от 24.12.2021; "
        f"разъяснение Роструда №15-2/В-1677 от 30.05.2022). Порядок и периодичность "
        f"инструктажей установлены {order_ref}. Подписи — собственноручные."
    )
    ws.merge_cells(start_row=row, start_column=1, end_row=row, end_column=NCOLS)
    c = ws.cell(row=row, column=1, value=footer_text)
    c.font = Font(name="Times New Roman", size=8, italic=True)
    c.alignment = Alignment(horizontal="left", vertical="top", wrap_text=True)

    widths = {"A": 8, "B": 7, "C": 12, "D": 30, "E": 12, "F": 20, "G": 22, "H": 22, "I": 14, "J": 14}
    for col_letter, w in widths.items():
        ws.column_dimensions[col_letter].width = w

    path = f"{output_dir}/journal_{instruction_type.value}_{datetime.utcnow():%Y%m%d%H%M%S}.xlsx"
    wb.save(path)
    return path


# ВРЕМЕННО (тест, 2026-07): очистка всех записей инструктажей, чтобы можно было
# заново проверить автозаполнение/допечатку с чистого листа. УДАЛИТЬ эту функцию
# и кнопку в webforms.py, когда тестирование закончится — см. пометку там же.
def test_clear_all_instructions(session: Session) -> int:
    count = session.query(Instruction).count()
    session.query(Instruction).delete()
    session.commit()
    return count


def get_due_instructions(session: Session, days_before: int = 7) -> list[dict]:
    """Повторные инструктажи, у которых next_due_date приближается или уже
    прошёл — для напоминания, по аналогии с get_rotation_reminders в tabel.py."""
    today = date.today()
    rows = (
        session.query(Instruction)
        .filter(Instruction.next_due_date.isnot(None))
        .all()
    )
    result = []
    for instr in rows:
        delta = (instr.next_due_date - today).days
        if delta <= days_before:
            emp = session.get(Employee, instr.employee_id)
            if emp is not None:
                result.append({
                    "employee_id": emp.id, "name": emp.full_name,
                    "due_date": instr.next_due_date, "overdue": delta < 0,
                    "instruction_id": instr.id,
                })
    return result


# ================= Удостоверения (корочки) =================

def create_certificate(
    session: Session, employee_id: str, profession: str,
    issued_by_org: str | None = None, issue_date: date | None = None,
    expiry_date: date | None = None, scan_key: str | None = None,
) -> Certificate:
    cert = Certificate(
        employee_id=employee_id,
        profession=profession,
        issued_by_org=issued_by_org,
        issue_date=issue_date,
        expiry_date=expiry_date,
        scan_key=scan_key,
    )
    session.add(cert)
    session.commit()
    return cert


def set_certificate_scan_key(session: Session, certificate_id: str, scan_key: str) -> None:
    cert = session.get(Certificate, certificate_id)
    if cert is not None:
        cert.scan_key = scan_key
        session.add(cert)
        session.commit()


def get_certificates_for_employee(session: Session, employee_id: str) -> list[Certificate]:
    return (
        session.query(Certificate)
        .filter_by(employee_id=employee_id)
        .order_by(Certificate.expiry_date.is_(None), Certificate.expiry_date)
        .all()
    )


def certificate_status(cert: Certificate, today: date | None = None) -> str:
    """'active' | 'expiring_soon' | 'expired' | 'no_expiry' — для бейджей в UI."""
    if cert.expiry_date is None:
        return "no_expiry"
    today = today or date.today()
    delta = (cert.expiry_date - today).days
    if delta < 0:
        return "expired"
    if delta <= CERTIFICATE_EXPIRY_WARNING_DAYS:
        return "expiring_soon"
    return "active"


def get_expiring_certificates(session: Session, days_before: int = CERTIFICATE_EXPIRY_WARNING_DAYS) -> list[dict]:
    """Удостоверения, срок которых истекает в пределах days_before дней (включая
    уже просроченные) — для напоминаний, та же схема, что и get_due_instructions."""
    today = date.today()
    rows = (
        session.query(Certificate)
        .filter(Certificate.expiry_date.isnot(None))
        .all()
    )
    result = []
    for cert in rows:
        delta = (cert.expiry_date - today).days
        if delta <= days_before:
            emp = session.get(Employee, cert.employee_id)
            if emp is not None:
                result.append({
                    "employee_id": emp.id, "name": emp.full_name,
                    "profession": cert.profession, "expiry_date": cert.expiry_date,
                    "overdue": delta < 0, "certificate_id": cert.id,
                })
    result.sort(key=lambda r: r["expiry_date"])
    return result


# ================= Печатный бланк наряда-допуска =================
# Печатный (не рукописный) наряд-допуск разрешён действующими Правилами по ОТ —
# "документ можно составлять на компьютере или от руки" (проверено, см. журнал
# патчей/обсуждение). Единственное, что должно остаться "живым" — подписи на
# распечатанном экземпляре, сам текст печатать можно.
#
# [Предполагаю] Это ОБЩИЙ бланк, не учитывает разницу форм по видам работ
# (электроустановки — приложение №7 к Приказу №903н, высотные — приложение №2
# и т.д. — формы РАЗНЫЕ и менять их содержание нельзя). Пока WorkOrder не имеет
# поля work_type, бланк один на все случаи — годится для общих/ремонтных работ,
# где унифицированной формы законом не установлено. Для регламентированных видов
# (электро-, высотные, огневые) нужен отдельный шаблон под конкретное приложение
# правил — это следующий шаг, не делать вид, что текущий бланк их закрывает.

def generate_work_order_docx(work_order: WorkOrder, members: list[WorkOrderMember],
                              org_name: str, output_dir: str = "/tmp") -> str:
    """
    2026-07: переписано по образцу реального бланка ООО «Промстроймонтаж»
    (наряд №25 на работы на высоте, фото прислал пользователь) — см. подробный
    docstring WorkOrder в models.py про структурные поправки. Разделы идут в
    том же порядке, что и в реальном бланке, для узнаваемости.
    """
    doc = Document()
    _set_a4(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(f"НАРЯД-ДОПУСК № {work_order.number}")
    run.bold = True
    run.font.size = Pt(16)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run("на производство работ повышенной опасности").italic = True

    doc.add_paragraph(f"Организация: {org_name}")
    if work_order.subdivision:
        doc.add_paragraph(f"Подразделение: {work_order.subdivision}")
    doc.add_paragraph(f"Выдан: «{work_order.created_at.strftime('%d')}» "
                       f"{work_order.created_at.strftime('%m.%Y')} г.")
    doc.add_paragraph(
        f"Действителен до: «{work_order.valid_to.strftime('%d')}» "
        f"{work_order.valid_to.strftime('%m.%Y')} г."
    )
    doc.add_paragraph(
        f"Ответственному руководителю работ: "
        f"{work_order.responsible_supervisor.full_name if work_order.responsible_supervisor else '—'}"
    )
    doc.add_paragraph(
        f"Ответственному исполнителю работ: "
        f"{work_order.responsible_executor.full_name if work_order.responsible_executor else '—'}"
    )
    doc.add_paragraph()
    doc.add_paragraph(f"На выполнение работ: {work_order.work_description}")
    doc.add_paragraph(f"Место выполнения работ: {work_order.location}")

    # Материалы/инструменты/приспособления/спецтехника — только если заполнены
    # (необязательные поля, не все наряды их требуют).
    for label, value in [
        ("Материалы", work_order.materials),
        ("Инструменты", work_order.tools),
        ("Приспособления", work_order.equipment),
        ("Спецтехника", work_order.special_machinery),
    ]:
        if value:
            doc.add_paragraph(f"{label}: {value}")

    if work_order.technological_card_ref:
        doc.add_paragraph()
        doc.add_paragraph(
            f"Работы производить в соответствии с требованиями технологической карты "
            f"({work_order.technological_card_ref})."
        )

    if work_order.safety_systems:
        doc.add_paragraph()
        doc.add_paragraph("Системы обеспечения безопасности работ:").bold = True
        doc.add_paragraph(work_order.safety_systems)

    if work_order.special_conditions:
        doc.add_paragraph()
        doc.add_paragraph("Особые условия проведения работ:").bold = True
        doc.add_paragraph(work_order.special_conditions)

    doc.add_paragraph()
    doc.add_paragraph("Состав исполнителей работ (члены бригады):").bold = True
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "№"
    hdr[1].text = "Фамилия, имя, отчество"
    hdr[2].text = "С условиями работ ознакомил, инструктаж провёл (подпись)"
    for i, member in enumerate(members, start=1):
        row = table.add_row().cells
        row[0].text = str(i)
        row[1].text = member.employee.full_name if member.employee else "?"
        row[2].text = ""  # печатный бланк — подпись ставится от руки на распечатке

    # Ежедневный допуск к работе — по одной строке на каждый день периода действия
    # наряда (см. WorkOrderDailyAdmission). Пустые графы — заполняются от руки/по
    # факту при подтверждении в системе (record_daily_briefing/completion).
    doc.add_paragraph()
    doc.add_paragraph("Ежедневный допуск к работе:").bold = True
    daily_table = doc.add_table(rows=1, cols=4)
    daily_table.style = "Table Grid"
    dhdr = daily_table.rows[0].cells
    dhdr[0].text = "Дата"
    dhdr[1].text = "Целевой инструктаж выдан (дата, время, подпись руководителя)"
    dhdr[2].text = "Работа закончена, место убрано (дата, время, подпись исполнителя)"
    dhdr[3].text = ""
    day = work_order.valid_from
    while day <= work_order.valid_to:
        drow = daily_table.add_row().cells
        drow[0].text = day.strftime("%d.%m.%Y")
        drow[1].text = ""
        drow[2].text = ""
        drow[3].text = ""
        day += timedelta(days=1)

    doc.add_paragraph()
    doc.add_paragraph("Подписи:")
    doc.add_paragraph("Наряд выдал: _________________________  (подпись, расшифровка)")
    doc.add_paragraph("Ответственный руководитель работ: _________________________  (подпись, расшифровка)")
    doc.add_paragraph("Ответственный исполнитель работ: _________________________  (подпись, расшифровка)")

    path = f"{output_dir}/naryad_{work_order.number}_{work_order.id[:8]}.docx"
    doc.save(path)
    return path


# ================= Наряд-допуск на работы на ВЫСОТЕ (782н, Приложение № 2) =================
# Отдельный бланк под регламентированную форму (см. замечание перед общим генератором выше).
# Наполняется из справочника типовых работ (WorkType) через WorkOrder.work_type_id.
# Общий generate_work_order_docx остаётся для нерегламентированных работ; печать ветвится
# по наличию work_type_id (см. webforms.py).

_HWO_DASH = "—"
_HWO_START_TIME = "08:00"
_HWO_END_TIME = "19:00"
_HWO_MAX_DAYS = 15  # 782н п. 65: срок действия ≤ 15 календарных дней (+1 продление ≤15)

_HWO_PREP_MEASURES = [
    "Оформить наряд-допуск на работы повышенной опасности с обязательным указанием: "
    "ответственного исполнителя работ; ответственного руководителя работ; место выполнения "
    "работ на высоте находится в зоне прямой видимости ответственного исполнителя и/или "
    "ответственного руководителя работ.",
    "Ознакомление и обсуждение Плана производства работ (технологической карты) с "
    "ответственным руководителем работ, ответственным исполнителем работ, исполнителями работ.",
    "Разъяснение ответственным руководителем работ специфических обязанностей и процедур всем "
    "работникам, соблюдение правил безопасности.",
    "Работники, впервые допускаемые к работам на высоте, должны обладать практическими "
    "навыками применения оборудования и оказания первой помощи, применения СИЗ, их осмотром "
    "до и после использования.",
    "Средства коллективной и индивидуальной защиты должны использоваться по назначению в "
    "соответствии с требованиями инструкций изготовителя и нормативной технической документации.",
]


def _hwo_group_num(value):
    """Из строки группы («2-я гр. по безопасности работ на высоте») достаёт число 2."""
    if not value:
        return None
    digits = ""
    for ch in str(value):
        if ch.isdigit():
            digits += ch
        elif digits:
            break
    return int(digits) if digits else None


def _hwo_split_systems(raw):
    """safety_systems (одно текстовое поле) -> 3 значения строк таблицы систем.
    < 3 строк -> недостающие прочерком; > 3 -> лишние склеиваются в последнюю."""
    lines = [ln.strip() for ln in (raw or "").splitlines() if ln.strip()]
    vals = []
    for i in range(3):
        if i < len(lines):
            vals.append(" ".join(lines[2:]) if (i == 2 and len(lines) > 3) else lines[i])
        else:
            vals.append(_HWO_DASH)
    return vals


def check_work_order_problems(work_order: WorkOrder, members=None) -> list[str]:
    """Нарушения наряда-допуска на высоте (782н + правила проекта). Пустой список = ок.
    members можно передать явно (как их видит вызывающий код); иначе берутся из work_order.
    Документ строится всегда (черновик); блокировать выпуск по этому списку — задача
    вызывающего кода (webforms.py)."""
    problems: list[str] = []
    sup = getattr(work_order, "responsible_supervisor", None)
    ex = getattr(work_order, "responsible_executor", None)

    if _hwo_group_num(getattr(sup, "height_safety_group", None)) != 3:
        problems.append(
            "Ответственный руководитель работ должен быть 3-й группы по безопасности работ на "
            f"высоте (сейчас: {getattr(sup, 'height_safety_group', None) or _HWO_DASH})."
        )
    ex_g = _hwo_group_num(getattr(ex, "height_safety_group", None))
    if ex_g is None or ex_g < 2:
        problems.append(
            "Ответственный исполнитель работ должен быть не ниже 2-й группы "
            f"(сейчас: {getattr(ex, 'height_safety_group', None) or _HWO_DASH})."
        )

    members = list(members if members is not None else (getattr(work_order, "members", None) or []))
    if not members:
        problems.append("Состав бригады пуст — нельзя выпустить наряд без исполнителей.")
    for m in members:
        emp = getattr(m, "employee", None)
        name = getattr(emp, "full_name", None)
        if not emp or not name:
            problems.append("В бригаде есть член без привязанного сотрудника (пустая строка).")
            continue
        g = _hwo_group_num(getattr(emp, "height_safety_group", None))
        if g is None or g < 2:
            problems.append(f"У члена бригады «{name}» не указана группа по высоте (нужна ≥2-й).")

    try:
        span = (work_order.valid_to - work_order.valid_from).days + 1
        if span > _HWO_MAX_DAYS:
            problems.append(f"Срок действия наряда {span} дн. превышает 15 календарных дней (782н).")
        if span < 1:
            problems.append("Дата окончания раньше даты начала.")
    except Exception:
        problems.append("Не заданы корректные даты периода.")

    rescue = _hwo_split_systems(getattr(work_order, "safety_systems", None))[2].lower()
    wt = getattr(work_order, "work_type", None)
    if not (work_order.safety_systems or "").strip() and wt is not None:
        rescue = (getattr(wt, "sys_rescue", None) or "").lower()
    if rescue and rescue != _HWO_DASH and ("привяз" in rescue or "строп" in rescue or "фал" in rescue):
        problems.append(
            "Строка «Эвакуационные и спасательные системы» указывает страховочную привязь/строп/"
            "фал — нужно реальное средство спасения (например, автогидроподъёмник)."
        )
    return problems


def generate_height_work_order_docx(work_order: WorkOrder, members: list[WorkOrderMember],
                                     org_name: str, output_dir: str = "/tmp") -> str:
    """Полный бланк наряда-допуска на работы на высоте (Приложение № 2 к Правилам 782н).
    Тексты, зависящие от вида работ (содержание, условия, ОВПФ, 3 системы, раздел 3, нормы),
    берутся из связанного WorkType; собственное поле наряда — в приоритете, справочник —
    запасной источник. Время 08:00/19:00 константами (v1). Возвращает путь к docx."""
    doc = Document()
    _set_a4(doc)

    def _p(text="", *, bold=False, italic=False, center=False, size=11):
        p = doc.add_paragraph()
        if center:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        r.bold = bold
        r.italic = italic
        r.font.size = Pt(size)
        return p

    def _label(label, value):
        p = doc.add_paragraph()
        r = p.add_run(label); r.bold = True; r.font.size = Pt(11)
        r2 = p.add_run(value if value not in (None, "") else _HWO_DASH); r2.font.size = Pt(11)
        return p

    def _cell(cell, text, *, bold=False, size=9, center=False):
        cell.text = ""
        para = cell.paragraphs[0]
        if center:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = para.add_run("" if text in (None, "") else str(text))
        r.bold = bold; r.font.size = Pt(size)

    def _grid(headers, rows):
        t = doc.add_table(rows=1, cols=len(headers))
        t.style = "Table Grid"
        for i, h in enumerate(headers):
            _cell(t.rows[0].cells[i], h, bold=True, size=9, center=True)
        for row in rows:
            cells = t.add_row().cells
            for i, val in enumerate(row):
                _cell(cells[i], val, size=9)
        return t

    sup = getattr(work_order, "responsible_supervisor", None)
    ex = getattr(work_order, "responsible_executor", None)
    sup_name = getattr(sup, "full_name", None) or "____________"
    ex_name = getattr(ex, "full_name", None) or "____________"
    subdivision = work_order.subdivision or "Мурманск"

    wt = getattr(work_order, "work_type", None)
    work_name = work_order.work_description or getattr(wt, "name", None)
    content_val = getattr(wt, "content", None) or work_order.work_description
    conditions_val = work_order.special_conditions or getattr(wt, "conditions", None)
    hazards_val = getattr(work_order, "hazards", None) or getattr(wt, "hazards", None)
    norms_val = getattr(wt, "norms", None)
    # Пункт 1: собственное поле наряда в приоритете, иначе — типовое из справочника.
    materials_val = work_order.materials or getattr(wt, "materials", None)
    tools_val = work_order.tools or getattr(wt, "tools", None)
    equipment_val = work_order.equipment or getattr(wt, "equipment", None)

    _p("Приложение № 2 к Правилам по охране труда при работе на высоте", italic=True, size=9)
    _p("(Приказ Минтруда России от 16.11.2020 № 782н)", italic=True, size=9)
    _p()
    _p("УТВЕРЖДАЮ:", bold=True)
    _p(org_name)
    _p("_________________ / ____________")
    _p("«____» _____________ 20___ г.", italic=True)
    _p()
    _p(f"НАРЯД-ДОПУСК № {work_order.number or _HWO_DASH}", bold=True, center=True, size=13)
    _p("НА ПРОИЗВОДСТВО РАБОТ НА ВЫСОТЕ", bold=True, center=True)
    _p()

    _label("Организация: ", org_name)
    _label("Подразделение: ", subdivision)
    _label("Выдан ", f"{work_order.valid_from.strftime('%d.%m.%Y') if work_order.valid_from else _HWO_DASH} года")
    _label("Действителен до ", f"{work_order.valid_to.strftime('%d.%m.%Y') if work_order.valid_to else _HWO_DASH} года")
    _label("Ответственному руководителю работ: ", getattr(sup, "full_name", None))
    _label("Ответственному исполнителю (производителю) работ: ", getattr(ex, "full_name", None))
    _label("На выполнение работ: ", work_name)

    _p()
    _p("Состав исполнителей работ (члены бригады):")
    member_rows = []
    for i, m in enumerate(members or [], 1):
        emp = getattr(m, "employee", None)
        nm = getattr(emp, "full_name", None) or _HWO_DASH
        pos = getattr(emp, "position", None) or ""
        grp = getattr(emp, "height_safety_group", None) or ""
        pos_grp = ", ".join([x for x in (pos, grp) if x]) or _HWO_DASH
        # Вариант Б: печатаем ФИО ответственного руководителя в графу «Инструктаж
        # провёл» (он проводит целевой инструктаж по п.8 бланка 782н); подпись он
        # ставит от руки. «Ознакомлен» оставляем пустым — там подпись работника.
        member_rows.append([str(i), nm, pos_grp, sup_name, ""])
    if not member_rows:
        member_rows = [["", _HWO_DASH, _HWO_DASH, sup_name, ""]]
    _grid(["№", "Фамилия, имя, отчество", "Должность (разряд)",
           "Инструктаж провёл (подпись)", "Ознакомлен (подпись)"], member_rows)

    _p()
    _label("Место выполнения работ: ", work_order.location)
    _label("Содержание работ: ", content_val)
    _label("Условия проведения работ: ", conditions_val)
    _label("Опасные и вредные производственные факторы, которые действуют или могут возникнуть "
           "в местах выполнения работ: ", hazards_val)
    _label("Начало работ: ", f"{_HWO_START_TIME} {work_order.valid_from.strftime('%d.%m.%Y') if work_order.valid_from else _HWO_DASH}")
    _label("Окончание работ: ", f"{_HWO_END_TIME} {work_order.valid_to.strftime('%d.%m.%Y') if work_order.valid_to else _HWO_DASH}")
    if norms_val:
        _label("Нормативные основания: ", norms_val)

    _p()
    _p("Системы обеспечения безопасности работ на высоте:", bold=True)
    if (work_order.safety_systems or "").strip():
        ss = _hwo_split_systems(work_order.safety_systems)
    elif wt is not None:
        ss = [getattr(wt, "sys_restraint", None) or _HWO_DASH,
              getattr(wt, "sys_fall_arrest", None) or _HWO_DASH,
              getattr(wt, "sys_rescue", None) or _HWO_DASH]
    else:
        ss = [_HWO_DASH, _HWO_DASH, _HWO_DASH]
    _grid(["Системы обеспечения безопасности", "Состав системы"],
          [["Удерживающие системы", ss[0]],
           ["Страховочные системы", ss[1]],
           ["Эвакуационные и спасательные системы", ss[2]]])

    _p()
    _p("1. Необходимые для производства работ:", bold=True)
    for lbl, val in [("Материалы: ", materials_val), ("Инструмент: ", tools_val),
                     ("Приспособления: ", equipment_val),
                     ("Спецтехника: ", work_order.special_machinery),
                     ("Шифр ТК: ", work_order.technological_card_ref)]:
        if val:
            _label(lbl, val)

    _p()
    _p("2. До начала работ следует выполнить следующие мероприятия:", bold=True)
    _grid(["Наименование мероприятия", "Срок выполнения", "Ответственный исполнитель"],
          [[m, "До начала работ", ""] for m in _HWO_PREP_MEASURES])

    _p()
    _p("3. В процессе производства работ необходимо выполнить следующие мероприятия:", bold=True)
    proc_lines = [ln.strip() for ln in (getattr(wt, "process_measures", None) or "").splitlines() if ln.strip()]
    proc_rows = ([[m, "Постоянно в процессе работ", ex_name] for m in proc_lines]
                 if proc_lines else [["", "", ""] for _ in range(3)])
    _grid(["Наименование мероприятия", "Срок выполнения", "Ответственный исполнитель"], proc_rows)

    _p()
    _p("4. Особые условия проведения работ:", bold=True)
    _grid(["Наименование условий", "Срок выполнения", "Ответственный исполнитель"],
          [["", "", ""] for _ in range(2)])

    _p()
    _p("Отдельные указания: _______________________________________________")
    _p(f"Наряд выдал: ______________ (дата, время)   Подпись: ____________ / {org_name}")
    _p("Наряд продлил: ____________ (дата, время)   Подпись: ____________ / ____________")

    _p()
    _p("5. Разрешение на подготовку рабочих мест и на допуск к выполнению работ:", bold=True)
    _grid(["Разрешение на подготовку и допуск получил", "Дата, время", "Подпись"], [["", "", ""]])
    _p("Рабочие места подготовлены. Ответственный руководитель работ: ____________ / " + sup_name)

    _p()
    _p("6. Ежедневный допуск к работе и время её окончания:", bold=True)
    day_rows = []
    try:
        d = work_order.valid_from
        while d <= work_order.valid_to:
            day_rows.append([d.strftime("%d.%m.%Y"), "", ""])
            d += timedelta(days=1)
    except Exception:
        pass
    if not day_rows:
        day_rows = [["", "", ""]]
    _grid(["Дата", "Бригада получила целевой инструктаж и допущена (дата, время, подпись)",
           "Работа закончена, бригада удалена (дата, время, подпись)"], day_rows)

    _p()
    _p("7. Изменения в составе бригады:", bold=True)
    change_rows = []
    for ch in getattr(work_order, "member_changes", None) or []:
        nm = getattr(getattr(ch, "employee", None), "full_name", None) or _HWO_DASH
        added = nm if ch.change_type == MemberChangeType.ADDED else ""
        removed = nm if ch.change_type == MemberChangeType.REMOVED else ""
        when = ch.changed_at.strftime("%d.%m.%Y %H:%M") if ch.changed_at else ""
        change_rows.append([added, removed, when, ch.ordered_by or ""])
    if not change_rows:
        change_rows = [["", "", "", ""] for _ in range(3)]  # пустые строки под ручную запись
    _grid(["Введён в состав (ФИО)", "Выведен из состава (ФИО)", "Дата, время", "Разрешил (ФИО, подпись)"],
          change_rows)

    _p()
    _p("8. Регистрация целевого инструктажа при первичном допуске:", bold=True)
    _p("Инструктаж провёл: ____________ / " + sup_name)
    _p("Инструктаж прошёл: ____________ / ____________")
    _p("Лицо, выдавшее наряд: ____________ / " + org_name)
    _p("Ответственный руководитель работ: ____________ / " + sup_name)
    _p("Ответственный исполнитель: ____________ / " + ex_name)

    _p()
    _p("9. Письменное разрешение (акт-допуск) действующего предприятия на производство работ "
       "имеется. Мероприятия по безопасности согласованы:", bold=True)
    _p("__________________________________________ (должность, ФИО, подпись)")

    _p()
    _p("10. Рабочее место и условия труда проверены. Мероприятия по безопасности выполнены. "
       "Разрешаю приступить к выполнению работ:", bold=True)
    _p("_______________________ (дата, подпись) ____________________ (ФИО)")
    _p("Наряд-допуск продлён до: ______________ (дата, подпись) ____________ (ФИО)")

    _p()
    _p("11. Работа выполнена в полном объёме. Материалы, инструмент, приспособления убраны. "
       "Члены бригады выведены.", bold=True)
    _p("Ответственный исполнитель (производитель) работ: ____________ (дата, подпись)")
    _p("Наряд-допуск закрыт.")
    _p("Ответственный руководитель работ: ____________ (дата, подпись)     "
       "Лицо, выдавшее наряд-допуск: ____________ (дата, подпись)")

    path = f"{output_dir}/naryad_vysota_{work_order.number}_{work_order.id[:8]}.docx"
    doc.save(path)
    return path


# ================= Реестр приказов =================

def create_order(session: Session, number: str, order_date: date, topic: str,
                  category: OrderCategory = OrderCategory.OTHER,
                  note: str | None = None, order_key: str | None = None) -> InternalOrder:
    order = InternalOrder(number=number, order_date=order_date, topic=topic,
                           category=category, note=note, order_key=order_key)
    session.add(order)
    session.commit()
    return order


def get_order_by_key(session: Session, order_key: str) -> "InternalOrder | None":
    """Запись реестра, привязанная к шаблону приказа ОТ (по order_key). Нужна, чтобы генерация
    не плодила дубли: если приказ уже заведён — переиспользуем запись, а не создаём вторую."""
    return (
        session.query(InternalOrder)
        .filter_by(order_key=order_key)
        .order_by(InternalOrder.created_at.desc())
        .first()
    )


def get_ot_orders_status(session: Session) -> list[dict]:
    """Контроль обязательных приказов по ОТ (все из OT_ORDER_KEYS). Для каждого — статус:
    'missing'   — не сгенерирован (нет записи в реестре);
    'no_scan'   — заведён, но скан не загружен (не распечатан/не подписан) — на это и контроль;
    'ready'     — есть запись со сканом.
    Возвращает список по порядку OT_ORDER_KEYS для показа в списке приказов."""
    result = []
    for key in OT_ORDER_KEYS:
        category, topic, _preamble, _points = OT_ORDERS[key]
        rec = get_order_by_key(session, key)
        if rec is None:
            status = "missing"
        elif not rec.scan_key:
            status = "no_scan"
        else:
            status = "ready"
        result.append({
            "key": key, "topic": topic, "category": category,
            "section": OT_SECTIONS.get(category, ""), "status": status,
            "record": rec,
        })
    return result


def get_orders(session: Session) -> list[InternalOrder]:
    return session.query(InternalOrder).order_by(InternalOrder.order_date.desc()).all()


def get_orders_by_category(session: Session, category: OrderCategory) -> list[InternalOrder]:
    return (
        session.query(InternalOrder)
        .filter_by(category=category)
        .order_by(InternalOrder.order_date.desc())
        .all()
    )


def set_order_scan_key(session: Session, order_id: str, scan_key: str) -> None:
    order = session.get(InternalOrder, order_id)
    if order is not None:
        order.scan_key = scan_key
        session.add(order)
        session.commit()


def get_latest_order_ref(session: Session) -> str:
    """
    Ссылка на актуальный приказ для футеров печатных бланков (наряд-допуск,
    журналы инструктажа) — INTERNAL_ORDER_REF. Берёт последний по дате приказ
    из реестра. Если реестр пуст — явная заглушка, чтобы не выглядело как
    настоящая ссылка на несуществующий документ (см. договорённость)."""
    orders = get_orders(session)
    if not orders:
        return "[Приказ не издан — заполнить в разделе «Приказы»]"
    latest = orders[0]
    return f"Приказ № {latest.number} от {latest.order_date.strftime('%d.%m.%Y')}"


# ================= Генератор приказов по охране труда =================
# Один генератор + справочник приказов (каждый приказ — данные: раздел, тема, основание,
# пункты). Реквизиты ИП Буц вшиты по умолчанию (согласовано). При смене подписанта/реквизитов
# менять здесь. Разделы соответствуют OrderCategory (LABOR_PROTECTION/TRAINING/HEIGHT_WORK/
# FIRE_SAFETY/ELECTRICAL).

OT_IP_NAME_FULL = "Индивидуальный предприниматель Буц Сергей Юрьевич"
OT_IP_NAME_SHORT = "ИП Буц С.Ю."
OT_IP_INN = "312608174376"
OT_IP_OGRNIP = "326508100306960"
OT_IP_SIGNER = "С.Ю. Буц"
OT_IP_RESPONSIBLE = "Буц С.Ю."
OT_WORK_PLACE = "с. Белокаменка, Мурманская область"

_OT_MONTHS = ['января', 'февраля', 'марта', 'апреля', 'мая', 'июня', 'июля',
              'августа', 'сентября', 'октября', 'ноября', 'декабря']

# Раздел (OrderCategory) -> человекочитаемое имя для UI
OT_SECTIONS = {
    OrderCategory.LABOR_PROTECTION: "Охрана труда",
    OrderCategory.TRAINING: "Обучение и инструктажи",
    OrderCategory.HEIGHT_WORK: "Работы на высоте",
    OrderCategory.FIRE_SAFETY: "Пожарная безопасность",
    OrderCategory.ELECTRICAL: "Электробезопасность",
}

_R = OT_IP_RESPONSIBLE
# key -> (категория, тема, основание, [пункты])
OT_ORDERS = {
    "suot": (OrderCategory.LABOR_PROTECTION,
        "Об утверждении Положения о системе управления охраной труда",
        "В соответствии со статьёй 217 Трудового кодекса Российской Федерации, приказом "
        "Минтруда России от 29.10.2021 № 776н «Об утверждении Примерного положения о системе "
        "управления охраной труда», в целях организации системы управления охраной труда,",
        ["Утвердить и ввести в действие Положение о системе управления охраной труда (СУОТ) "
         "согласно приложению к настоящему приказу.",
         f"Ответственному за охрану труда {_R} ознакомить работников с Положением о СУОТ под роспись.",
         "Настоящий приказ довести до сведения всех работников.",
         "Контроль за исполнением настоящего приказа оставляю за собой."]),
    "responsible_ot": (OrderCategory.LABOR_PROTECTION,
        "О назначении ответственного за организацию работы по охране труда",
        "В соответствии со статьёй 217 Трудового кодекса Российской Федерации, в целях организации "
        "работы по охране труда и обеспечения безопасных условий труда работников,",
        [f"Назначить ответственным за организацию работы по охране труда {_R}.",
         "Возложить на ответственного обязанности по обеспечению соблюдения требований охраны труда, "
         "проведению инструктажей, контролю за состоянием условий труда и учёту микротравм.",
         "Ответственному руководствоваться в работе Положением о СУОТ и требованиями охраны труда.",
         "Контроль за исполнением настоящего приказа оставляю за собой."]),
    "instructions_ot": (OrderCategory.LABOR_PROTECTION,
        "Об утверждении инструкций по охране труда",
        "В соответствии со статьёй 214 Трудового кодекса Российской Федерации, приказом Минтруда "
        "России от 29.10.2021 № 772н «Об утверждении Основных требований к порядку разработки и "
        "содержания правил и инструкций по охране труда»,",
        ["Утвердить и ввести в действие инструкции по охране труда по профессиям и видам работ "
         "согласно перечню (приложение).",
         f"Ответственному за охрану труда {_R} ознакомить работников с инструкциями под роспись и "
         "обеспечить их пересмотр не реже одного раза в 5 лет.",
         "Контроль за исполнением настоящего приказа оставляю за собой."]),
    "responsible_siz": (OrderCategory.LABOR_PROTECTION,
        "О назначении ответственного за обеспечение работников СИЗ",
        "В соответствии со статьёй 221 Трудового кодекса Российской Федерации, приказом Минтруда "
        "России от 29.10.2021 № 766н, в целях обеспечения работников средствами индивидуальной защиты,",
        [f"Назначить ответственным за выдачу, учёт и хранение средств индивидуальной защиты (СИЗ) и "
         f"смывающих средств {_R}.",
         "Обеспечить выдачу СИЗ по установленным нормам с ведением личных карточек учёта.",
         "Контроль за исполнением настоящего приказа оставляю за собой."]),
    "training": (OrderCategory.TRAINING,
        "Об организации обучения и проверки знания требований охраны труда",
        "В соответствии со статьёй 219 Трудового кодекса Российской Федерации, постановлением "
        "Правительства Российской Федерации от 24.12.2021 № 2464 «О порядке обучения по охране "
        "труда и проверки знания требований охраны труда»,",
        ["Организовать обучение работников требованиям охраны труда, оказанию первой помощи и "
         "применению СИЗ с последующей проверкой знаний.",
         "Создать комиссию по проверке знания требований охраны труда в составе не менее трёх "
         "человек, прошедших обучение в установленном порядке.",
         f"Ответственному за охрану труда {_R} обеспечить ведение программ обучения и протоколов "
         "проверки знаний.",
         "Контроль за исполнением настоящего приказа оставляю за собой."]),
    "briefings": (OrderCategory.TRAINING,
        "Об организации проведения инструктажей по охране труда",
        "В соответствии со статьёй 219 Трудового кодекса Российской Федерации, постановлением "
        "Правительства Российской Федерации от 24.12.2021 № 2464, ГОСТ 12.0.004-2015,",
        ["Установить обязательное проведение вводного, первичного, повторного, внепланового и "
         "целевого инструктажей по охране труда в установленные сроки.",
         f"Назначить {_R} ответственным за проведение вводного инструктажа и ведение журналов "
         "регистрации инструктажей.",
         "Повторный инструктаж проводить не реже одного раза в 6 месяцев.",
         "Контроль за исполнением настоящего приказа оставляю за собой."]),
    "admission_work": (OrderCategory.TRAINING,
        "О допуске работников к самостоятельной работе",
        "В соответствии со статьёй 76 Трудового кодекса Российской Федерации, постановлением "
        "Правительства Российской Федерации от 24.12.2021 № 2464, по итогам обучения, стажировки "
        "и проверки знаний требований охраны труда,",
        ["Допустить к самостоятельной работе работников, прошедших вводный и первичный инструктажи, "
         "стажировку на рабочем месте и проверку знания требований охраны труда.",
         "К работе не допускать лиц, не прошедших обучение, инструктаж и проверку знаний.",
         "Контроль за исполнением настоящего приказа оставляю за собой."]),
    "responsible_height": (OrderCategory.HEIGHT_WORK,
        "О назначении ответственного за организацию и безопасное проведение работ на высоте",
        "В соответствии с Правилами по охране труда при работе на высоте, утверждёнными приказом "
        "Минтруда России от 16.11.2020 № 782н, в целях обеспечения безопасного проведения работ на высоте,",
        [f"Назначить {_R} (3-я группа по безопасности работ на высоте) ответственным за организацию "
         "и безопасное проведение работ на высоте.",
         "Возложить на ответственного обязанности по организации выдачи нарядов-допусков, "
         "инструктажей и контролю применения систем обеспечения безопасности.",
         "Контроль за исполнением настоящего приказа оставляю за собой."]),
    "height_list": (OrderCategory.HEIGHT_WORK,
        "Об утверждении перечня работ на высоте, выполняемых по наряду-допуску",
        "В соответствии с пунктами 20–25 Правил по охране труда при работе на высоте (приказ "
        "Минтруда России от 16.11.2020 № 782н),",
        ["Утвердить перечень работ на высоте, выполняемых с оформлением наряда-допуска, согласно "
         "приложению.",
         "Работы на высоте, включённые в перечень, выполнять только при наличии оформленного "
         "наряда-допуска.",
         "Контроль за исполнением настоящего приказа оставляю за собой."]),
    "height_naryad": (OrderCategory.HEIGHT_WORK,
        "О назначении лиц, ответственных за выдачу нарядов-допусков на работы на высоте",
        "В соответствии с Правилами по охране труда при работе на высоте (приказ Минтруда России "
        "от 16.11.2020 № 782н),",
        [f"Назначить {_R} лицом, имеющим право выдавать наряды-допуски на производство работ на высоте.",
         "Установить, что наряд-допуск оформляется до начала работ и хранится в установленном порядке.",
         "Контроль за исполнением настоящего приказа оставляю за собой."]),
    "fire_responsible": (OrderCategory.FIRE_SAFETY,
        "О назначении ответственного за пожарную безопасность",
        "В соответствии со статьёй 37 Федерального закона от 21.12.1994 № 69-ФЗ «О пожарной "
        "безопасности», Правилами противопожарного режима в РФ (постановление Правительства РФ "
        "от 16.09.2020 № 1479),",
        [f"Назначить {_R}, прошедшего обучение мерам пожарной безопасности, ответственным за "
         "пожарную безопасность.",
         "Возложить на ответственного обеспечение соблюдения противопожарного режима, исправности "
         "первичных средств пожаротушения и проведение противопожарных инструктажей.",
         "Контроль за исполнением настоящего приказа оставляю за собой."]),
    "fire_regime": (OrderCategory.FIRE_SAFETY,
        "Об установлении противопожарного режима",
        "В соответствии с Правилами противопожарного режима в Российской Федерации, утверждёнными "
        "постановлением Правительства Российской Федерации от 16.09.2020 № 1479,",
        ["Установить на объектах противопожарный режим, определив места для курения, порядок "
         "уборки горючих отходов, обесточивания оборудования по окончании работ.",
         "Запретить проведение огневых работ без оформления наряда-допуска на огневые работы.",
         "Ответственному за пожарную безопасность провести противопожарный инструктаж работников.",
         "Контроль за исполнением настоящего приказа оставляю за собой."]),
    "electro_group1": (OrderCategory.ELECTRICAL,
        "О присвоении I группы по электробезопасности неэлектротехническому персоналу",
        "В соответствии с Правилами по охране труда при эксплуатации электроустановок (приказ "
        "Минтруда России от 15.12.2020 № 903н), Правилами технической эксплуатации электроустановок "
        "потребителей,",
        [f"Назначить {_R} ответственным за присвоение I группы по электробезопасности "
         "неэлектротехническому персоналу.",
         "Провести инструктаж неэлектротехническому персоналу с проверкой знаний и присвоением "
         "I группы по электробезопасности с ежегодным подтверждением.",
         "Контроль за исполнением настоящего приказа оставляю за собой."]),
}

# Порядок вывода в UI (по разделам)
OT_ORDER_KEYS = [
    "suot", "responsible_ot", "instructions_ot", "responsible_siz",
    "training", "briefings", "admission_work",
    "responsible_height", "height_list", "height_naryad",
    "fire_responsible", "fire_regime",
    "electro_group1",
]


def generate_ot_order_docx(order_key: str, number: str, order_date: date, output_dir: str) -> str:
    """Генерирует приказ по охране труда из справочника OT_ORDERS. Возвращает путь к .docx.
    Реквизиты ИП Буц вшиты (OT_IP_*). Один шаблон на все приказы, содержимое — из справочника."""
    if order_key not in OT_ORDERS:
        raise ValueError(f"Неизвестный приказ ОТ: {order_key}")
    category, topic, preamble, points = OT_ORDERS[order_key]

    doc = Document()
    sec = doc.sections[0]
    sec.page_width, sec.page_height = Mm(210), Mm(297)
    sec.left_margin = sec.right_margin = Mm(20)
    sec.top_margin = sec.bottom_margin = Mm(15)
    st = doc.styles["Normal"]
    st.font.name = "Times New Roman"
    st.font.size = Pt(12)

    def _op(text="", *, align=WD_ALIGN_PARAGRAPH.LEFT, bold=False, size=12, after=6):
        par = doc.add_paragraph()
        par.alignment = align
        par.paragraph_format.space_after = Pt(after)
        run = par.add_run(text)
        run.bold = bold
        run.font.size = Pt(size)
        return par

    _op(OT_IP_NAME_FULL, align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=13, after=2)
    _op(f"ИНН {OT_IP_INN}   ОГРНИП {OT_IP_OGRNIP}", align=WD_ALIGN_PARAGRAPH.CENTER, size=10, after=12)
    _op("ПРИКАЗ", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=14, after=2)
    _op(f"№ {number} от «{order_date.day:02d}» {_OT_MONTHS[order_date.month - 1]} {order_date.year} г.",
        align=WD_ALIGN_PARAGRAPH.CENTER, after=4)
    _op(OT_WORK_PLACE, align=WD_ALIGN_PARAGRAPH.CENTER, size=11, after=12)
    _op(topic, align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, after=12)
    _op(preamble, after=8)
    _op("ПРИКАЗЫВАЮ:", bold=True, after=8)
    for i, pt in enumerate(points, 1):
        _op(f"{i}. {pt}", after=6)
    _op("", after=14)
    par = doc.add_paragraph()
    par.paragraph_format.space_after = Pt(6)
    par.add_run(OT_IP_NAME_SHORT + " " * 18 + "_______________ / " + OT_IP_SIGNER + " /")
    _op("", after=12)
    _op("С приказом ознакомлен(ы):", size=11, after=10)
    _op("_______________ / _________________ /   «__» ________ " + str(order_date.year) + " г.", size=11)

    import os
    out_path = os.path.join(output_dir, f"prikaz_{order_key}.docx")
    doc.save(out_path)
    return out_path


# ================= Погода (Open-Meteo) =================
# Координаты объекта (Белокаменка, западный берег Кольского залива).
WEATHER_LAT = 69.09
WEATHER_LON = 33.30
_WEATHER_URL = "https://api.open-meteo.com/v1/forecast"


def _weather_code_to_text(code: int) -> str:
    """Код погоды WMO → краткое описание (для ОЖР)."""
    if code == 0:
        return "ясно"
    if code in (1, 2, 3):
        return "переменная облачность"
    if code in (45, 48):
        return "туман"
    if 51 <= code <= 67:
        return "дождь"
    if 71 <= code <= 77:
        return "снег"
    if 80 <= code <= 82:
        return "ливень"
    if 95 <= code <= 99:
        return "гроза"
    return "н/д"


def update_weather(session: Session, past_days: int = 3, forecast_days: int = 1) -> int:
    """Тянет погоду из Open-Meteo и пишет/обновляет записи в таблице weather.
    По умолчанию берёт последние 3 дня + сегодня (чтобы закрыть пропуски, если
    планировщик не отработал). Возвращает число обновлённых дат. Плановый вызов —
    раз в день через APScheduler. Сеть: исходящий HTTPS к api.open-meteo.com
    (проверено — DPI на сервере не режет)."""
    import urllib.request
    import urllib.parse
    params = {
        "latitude": WEATHER_LAT,
        "longitude": WEATHER_LON,
        "daily": "temperature_2m_mean,weather_code",
        "timezone": "Europe/Moscow",
        "past_days": past_days,
        "forecast_days": forecast_days,
    }
    url = _WEATHER_URL + "?" + urllib.parse.urlencode(params)
    try:
        with urllib.request.urlopen(url, timeout=20) as resp:
            data = json.loads(resp.read().decode("utf-8"))
    except Exception as e:  # noqa: BLE001 — сеть/парсинг; не роняем планировщик
        log.warning("update_weather: не удалось получить погоду: %s", e)
        return 0
    daily = data.get("daily") or {}
    dates = daily.get("time") or []
    temps = daily.get("temperature_2m_mean") or []
    codes = daily.get("weather_code") or []
    updated = 0
    for d_str, t, c in zip(dates, temps, codes):
        if t is None:
            continue
        try:
            wd = datetime.strptime(d_str, "%Y-%m-%d").date()
        except ValueError:
            continue
        desc = _weather_code_to_text(int(c)) if c is not None else None
        row = session.query(Weather).filter_by(weather_date=wd).first()
        if row is None:
            session.add(Weather(weather_date=wd, temperature=round(float(t), 1), description=desc))
        else:
            row.temperature = round(float(t), 1)
            row.description = desc
        updated += 1
    session.commit()
    return updated
