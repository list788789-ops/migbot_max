"""
Генерация docx-документов для отправки сотруднику через бота:
- Согласие на обработку персональных данных (152-ФЗ)
- Направление на медицинский осмотр (медкомиссию, ст.13.3 115-ФЗ)

ВАЖНО: для согласия по 152-ФЗ обязательна идентификация оператора (работодателя) —
без неё документ не соответствует требованию "конкретности" (ч.1 ст.9 152-ФЗ).
Данные оператора берутся из переменных окружения — см. COMPANY_* ниже.
ПЕРЕД ПЕРВЫМ ИСПОЛЬЗОВАНИЕМ обязательно заполнить реальными данными компании:
плейсхолдеры в квадратных скобках не являются юридически значимыми и их наличие
в отправленном документе означает, что согласие не имеет силы.

С 1 сентября 2025 согласие на обработку ПД должно оформляться ОТДЕЛЬНЫМ документом,
не как часть трудового договора или другой формы — поэтому это отдельный docx,
а не пункт в существующем consent_texts.py (тот — для текста в чате, не для подписи).

2026-07: добавлена проверка обязательных полей сотрудника ПЕРЕД генерацией — раньше
пустые passport_series/passport_number/birth_date/address молча превращались в текст
"[не указано]" внутри готового документа, который уходит в клинику или сотруднику на
подпись. Теперь генератор поднимает ValueError с точным списком недостающих полей —
ВАЖНО: вызывающий код (bot.py, webforms.py) должен показывать текст ЭТОГО исключения
кадровику, а не общее "не удалось сгенерировать документ", иначе смысл проверки теряется.

2026-07: добавлен блок подписей "От Исполнителя / От Заказчика" в направление на
медосмотр — был в бумажном шаблоне (Приложение №1 к договору №176), но отсутствовал
в генераторе; документ обрывался на пункте 10.

2026-07: добавлен ТЕСТОВЫЙ режим TEST_ALLOW_MISSING_FIELDS — если включён, генератор
НЕ поднимает ValueError при незаполненных полях, а подставляет прочерк "—" и добавляет
явное предупреждение внутрь самого документа (чтобы черновик нельзя было спутать
с юридически валидным документом, если он случайно попадёт клинике или сотруднику).
Это временное послабление ТОЛЬКО для тестирования потока — флаг должен быть выключен
(или переменная удалена) до реальной работы с сотрудниками. Один флаг на оба генератора,
чтобы не потерять место, где включали, при отключении.

2026-07: у generate_consent_docx добавлен параметр require_fields (по умолчанию True).
require_fields=False — штатный режим ПЕЧАТИ бланка согласия под подпись: бланк можно
распечатать даже при незаполненных полях (кадровик печатает пустографку и заполняет от
руки/сотрудник подписывает). В этом режиме генератор НЕ поднимает ValueError, а просто
подставляет прочерки "—" на месте незаполненных полей, без каких-либо пометок в документе.
Это НЕ трогает остальные документы
(медосмотр, трудовой договор и т.д.) — только согласие и только по явному запросу
вызывающего кода (см. bot._send_consent_pdf / _handle_send_document для /send_consent_doc).
Веб-формы и прочие вызовы generate_consent_docx без параметра работают как раньше (строгий
режим), потому что по умолчанию require_fields=True.

Требуемые переменные окружения:
COMPANY_NAME — полное наименование юрлица-работодателя
COMPANY_INN — ИНН
COMPANY_LEGAL_ADDRESS — юридический адрес
HR_SIGNATORY_NAME — ФИО подписанта со стороны работодателя
HR_SIGNATORY_POSITION — должность подписанта
CLINIC_NAME — ПОЛНОЕ наименование медицинской организации (шапка бланка)
CLINIC_SHORT_NAME — короткое имя МО для блока подписи (напр. ГОАУЗ «МОМЦ»)
CLINIC_CONTRACT_NUMBER — номер договора с клиникой
CLINIC_CONTRACT_DATE — дата договора с клиникой, формат ДД.ММ.ГГГГ
CLINIC_CHIEF_DOCTOR_NAME — ФИО главного врача клиники (для блока подписи "От Исполнителя")
PAYER_NAME — заказчик услуги (напр. "ИП Буц С.Ю.") — используется в п.5/8 бланка
PAYER_SIGNATORY_NAME — ФИО подписанта со стороны заказчика для блока подписи
(напр. "С. Ю. Буц") — если не задано, используется PAYER_NAME
PAYER_PHONE — телефон заказчика
TEST_ALLOW_MISSING_FIELDS — "true"/"1" чтобы разрешить генерацию с прочерками вместо
незаполненных полей (ТОЛЬКО для теста, см. выше)
"""

import os
from datetime import date, datetime, timedelta, timezone

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from openpyxl import Workbook
from openpyxl.styles import Font

from models import Employee

MSK = timezone(timedelta(hours=3))  # Мурманская обл. — московское время, без перехода на летнее с 2014

COMPANY_NAME = os.environ.get("COMPANY_NAME", "[НЕ ЗАПОЛНЕНО — укажите наименование юрлица]")
COMPANY_INN = os.environ.get("COMPANY_INN", "[ИНН не указан]")
COMPANY_ADDRESS = os.environ.get("COMPANY_LEGAL_ADDRESS", "[юридический адрес не указан]")
HR_SIGNATORY_NAME = os.environ.get("HR_SIGNATORY_NAME", "[ФИО подписанта не указано]")
HR_SIGNATORY_POSITION = os.environ.get("HR_SIGNATORY_POSITION", "[должность не указана]")

# ТЕСТОВЫЙ флаг — см. заголовок файла. Читается один раз при импорте модуля;
# если меняешь переменную окружения на Railway, нужен рестарт сервиса, чтобы применилось.
TEST_ALLOW_MISSING_FIELDS = os.environ.get("TEST_ALLOW_MISSING_FIELDS", "false").strip().lower() in (
    "1", "true", "yes",
)

DASH = "—"

# Поля, обязательные для обоих документов — вынесены в константы, чтобы webforms.py/bot.py
# могли получить список отсутствующих полей ДО генерации (для баннера/сообщения), не только
# через перехват ValueError.
CONSENT_REQUIRED_FIELDS = {
    "birth_date": "дата рождения",
    "passport_series": "серия паспорта",
    "passport_number": "номер паспорта",
    "address": "адрес места пребывания",
}
MEDICAL_REFERRAL_REQUIRED_FIELDS = {
    "birth_date": "дата рождения",
    "passport_series": "серия паспорта",
    "passport_number": "номер паспорта",
    # address убран (2026-07): п.3 направления берёт SITE_ADDRESS (константа площадки),
    # employee.address в направление больше не идёт, поэтому его пустота не должна
    # блокировать генерацию. В согласии employee.address по-прежнему используется.
}


def _set_default_style(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)


def _passport_str(employee: Employee) -> str:
    passport = f"{employee.passport_series or ''} {employee.passport_number or ''}".strip()
    return passport or DASH


def _date_or_dash(d: date | None) -> str:
    return d.strftime("%d.%m.%Y") if d else DASH


def _text_or_dash(value: str | None) -> str:
    return value if value else DASH


def _missing_fields(employee: Employee, field_labels: dict[str, str]) -> list[str]:
    """Возвращает список человекочитаемых имён незаполненных полей, ничего не бросая.
    Использовать это, когда нужен просто список (баннер в UI, сообщение бота) —
    без побочного эффекта в виде исключения."""
    return [
        label for attr, label in field_labels.items()
        if getattr(employee, attr) in (None, "")
    ]


def check_consent_fields(employee: Employee) -> list[str]:
    """Список отсутствующих полей для согласия — вызывать из webforms.py/bot.py
    до генерации, если нужно показать баннер независимо от режима TEST_ALLOW_MISSING_FIELDS."""
    return _missing_fields(employee, CONSENT_REQUIRED_FIELDS)


def check_medical_referral_fields(employee: Employee) -> list[str]:
    """То же самое для направления на медосмотр."""
    return _missing_fields(employee, MEDICAL_REFERRAL_REQUIRED_FIELDS)


def _require_fields(employee: Employee, field_labels: dict[str, str]) -> list[str]:
    """Обычный режим: поднимает ValueError с точным списком недостающих полей, вместо
    того чтобы молча вставить в документ текст-плейсхолдер.

    Тестовый режим (TEST_ALLOW_MISSING_FIELDS=true): НЕ поднимает исключение, а
    возвращает список отсутствующих полей — вызывающий генератор обязан сам подставить
    прочерки в текст документа и добавить предупреждение (см. generate_*_docx ниже).
    """
    missing = _missing_fields(employee, field_labels)
    if missing and not TEST_ALLOW_MISSING_FIELDS:
        raise ValueError(
            f"Нельзя сгенерировать документ для {employee.full_name} — "
            f"не заполнены поля: {', '.join(missing)}. "
            f"Заполните их в карточке сотрудника перед генерацией."
        )
    return missing


def _add_test_warning_paragraph(doc: Document, missing: list[str]) -> None:
    """Явное предупреждение прямо в теле документа, если он сгенерирован с прочерками
    (тестовый режим TEST_ALLOW_MISSING_FIELDS или согласие в режиме require_fields=False).
    Цель — чтобы черновик нельзя было перепутать с юридически валидным документом, если
    он случайно уйдёт клинике или сотруднику на подпись."""
    warning = doc.add_paragraph()
    run = warning.add_run(
        "⚠ ЧЕРНОВИК — не заполнены поля: "
        + ", ".join(missing)
        + ". Документ не имеет юридической силы, пока эти поля не указаны в карточке "
        "сотрудника (или вписаны от руки) и документ не перегенерирован."
    )
    run.bold = True
    doc.add_paragraph()


def generate_consent_docx(
    employee: Employee,
    operator: str = "tsm",
    output_dir: str = "/tmp",
    require_fields: bool = True,
) -> str:
    """Согласие на обработку персональных данных — отдельный документ (152-ФЗ).
    operator: 'tsm' (ООО «ТРЕСТСТРОЙМОНТАЖ», по умолчанию) или 'ip' (ИП Буц С.Ю.).
    Реквизиты оператора — из PD_OPERATORS.

    require_fields (по умолчанию True): строгий режим — при незаполненных обязательных
    полях поднимается ValueError (через _require_fields), как раньше. Веб-формы и прочие
    вызовы без этого параметра работают как прежде.

    require_fields=False: режим ПЕЧАТИ бланка под подпись — бланк формируется даже при
    незаполненных полях (пустые поля становятся прочерками "—", без пометок), ValueError
    не поднимается. Нужен,
    чтобы кадровик мог распечатать согласие для любого сотрудника независимо от полноты
    карточки (см. bot._send_consent_pdf и /send_consent_doc). Затрагивает ТОЛЬКО согласие."""
    if require_fields:
        missing = _require_fields(employee, CONSENT_REQUIRED_FIELDS)
    else:
        # Печать бланка под подпись: пустые поля станут прочерками «—»,
        # ValueError не поднимаем. Никаких пометок в документ не добавляем.
        check_consent_fields(employee)

    doc = Document()
    _set_default_style(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("СОГЛАСИЕ\nна обработку персональных данных")
    run.bold = True
    run.font.size = Pt(14)

    doc.add_paragraph()

    birth = _date_or_dash(employee.birth_date)
    address = _text_or_dash(employee.address)

    op = PD_OPERATORS.get(operator, PD_OPERATORS[DEFAULT_PD_OPERATOR])
    body = (
        f"Я, {employee.full_name}, {birth} года рождения, "
        f"паспорт {_passport_str(employee)}, зарегистрированный(-ая) по адресу: {address}, "
        f"в соответствии со ст. 9 Федерального закона от 27.07.2006 № 152-ФЗ "
        f"«О персональных данных» даю согласие {op['name']} (ИНН {op['inn']}, "
        f"адрес: {op['address']}) (далее — Оператор) на обработку моих "
        f"персональных данных в целях исполнения трудового договора, ведения "
        f"миграционного учёта и исполнения обязанностей работодателя, "
        f"предусмотренных законодательством Российской Федерации о правовом "
        f"положении иностранных граждан."
    )
    doc.add_paragraph(body)

    doc.add_paragraph("Согласие даётся на обработку следующих персональных данных:")
    items = [
        "фамилия, имя, отчество; дата и место рождения; гражданство;",
        "паспортные данные, миграционная карта, данные о постановке на миграционный учёт;",
        "адрес регистрации и фактического проживания; контактный телефон;",
        "сведения о трудовом договоре, занимаемой должности;",
        "сведения о результатах медицинского осмотра, необходимые для допуска к работе;",
        "иные данные, прямо предусмотренные законодательством о миграционном учёте иностранных граждан.",
    ]
    for item in items:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_paragraph(
        "Согласие действует с момента подписания до его отзыва в письменной форме. "
        "Я проинформирован(-а) о праве отозвать настоящее согласие в любой момент, "
        "направив письменное заявление Оператору (ч. 2 ст. 9 152-ФЗ)."
    )

    doc.add_paragraph()
    doc.add_paragraph(f"Дата: {datetime.now(MSK).date().strftime('%d.%m.%Y')}")
    doc.add_paragraph(f"Подпись: _____________________ / {employee.full_name}")

    filename = f"consent_{operator}_{employee.id}.docx"
    path = os.path.join(output_dir, filename)
    doc.save(path)
    return path

# Реквизиты клиники и заказчика. Значения по умолчанию зашиты как fallback (по образцу
# направления «Пирогова», договор №176), но переменная окружения их переопределяет —
# при смене договора/главврача/телефона правь env в Railway, а не этот файл.
# CLINIC_NAME — ПОЛНОЕ имя МО (идёт в шапку "В ..."); CLINIC_SHORT_NAME — короткое,
# для блока подписи "От Исполнителя" (в образце там "ГОАУЗ «МОМЦ»", а не полное имя).
CLINIC_NAME = os.environ.get(
    "CLINIC_NAME",
    "Государственное областное автономное учреждение здравоохранения "
    "«Мурманский областной медицинский центр»",
)
CLINIC_SHORT_NAME = os.environ.get("CLINIC_SHORT_NAME", "ГОАУЗ «МОМЦ»")
CLINIC_CONTRACT_NUMBER = os.environ.get("CLINIC_CONTRACT_NUMBER", "176")
CLINIC_CONTRACT_DATE = os.environ.get("CLINIC_CONTRACT_DATE", "25.06.2026")
CLINIC_CHIEF_DOCTOR_NAME = os.environ.get("CLINIC_CHIEF_DOCTOR_NAME", "А.М. Амозов")
PAYER_NAME = os.environ.get("PAYER_NAME", "ИП Буц С.Ю.")
PAYER_SIGNATORY_NAME = os.environ.get("PAYER_SIGNATORY_NAME", "С. Ю. Буц")
PAYER_PHONE = os.environ.get("PAYER_PHONE", "+7 (985) 415-54-20")

# Адрес площадки — идёт в п.3 направления вместо employee.address (у всех вахтовиков
# один адрес проживания). Только для НАПРАВЛЕНИЯ; в согласии (152-ФЗ) остаётся личный
# employee.address. Переопределяется env, если площадка сменится.
SITE_ADDRESS = os.environ.get(
    "SITE_ADDRESS", "Мурманская обл., Кольский р-н, с. Белокаменка, зд. 1А, 184664"
)

# --- Реквизиты ООО «ТРЕСТСТРОЙМОНТАЖ» (работодатель по трудовому договору) ---
# ВАЖНО: это ДРУГОЕ юрлицо, не ИП Буц (тот — заказчик медуслуги). Не путать.
# Значения по образцу договора №0074 (Уристемов), приняты как факт без проверки.
# Все переопределяемы через env при смене реквизитов.
EMPLOYER_NAME_FULL = os.environ.get(
    "EMPLOYER_NAME_FULL", 'Общество с ограниченной ответственностью "ТРЕСТСТРОЙМОНТАЖ"'
)
EMPLOYER_NAME_SHORT = os.environ.get("EMPLOYER_NAME_SHORT", "ООО «ТРЕСТСТРОЙМОНТАЖ»")
EMPLOYER_INN = os.environ.get("EMPLOYER_INN", "5038107922")
EMPLOYER_KPP = os.environ.get("EMPLOYER_KPP", "770501001")
EMPLOYER_LEGAL_ADDRESS = os.environ.get(
    "EMPLOYER_LEGAL_ADDRESS",
    "115114, Город Москва, вн.тер.г. муниципальный округ Замоскворечье, "
    "наб Шлюзовая, д. 8, стр. 1, помещ. 3ВН",
)
EMPLOYER_ACTUAL_ADDRESS = os.environ.get("EMPLOYER_ACTUAL_ADDRESS", EMPLOYER_LEGAL_ADDRESS)
EMPLOYER_PHONE = os.environ.get("EMPLOYER_PHONE", "+7 (495) 147-82-79")
EMPLOYER_DIRECTOR_FULL = os.environ.get("EMPLOYER_DIRECTOR_FULL", "Железняка Валерия Александровича")

# Реестр операторов ПДн для согласия (152-ФЗ). Статичные реквизиты — правятся здесь под ревью git.
# Реквизиты ИП Буц взяты из договора СМР № ПСМ-ИПБ-16-00106 (раздел «Реквизиты Сторон»).
# ОГРН/ОГРНИП в согласии не печатаются по решению — только наименование, ИНН, адрес.
PD_OPERATORS = {
    "tsm": {
        "name": EMPLOYER_NAME_FULL,
        "inn": EMPLOYER_INN,
        "address": EMPLOYER_LEGAL_ADDRESS,
    },
    "ip": {
        "name": "Индивидуальный предприниматель Буц Сергей Юрьевич",
        "inn": "312608174376",
        "address": "141370, Московская обл., Сергиево-Посадский р-н, "
        "г. Хотьково, ул. Менделеева, д. 17, кв. 62",
    },
}
DEFAULT_PD_OPERATOR = "tsm"
EMPLOYER_DIRECTOR_SHORT = os.environ.get("EMPLOYER_DIRECTOR_SHORT", "Железняк В. А.")
# ОКВЭД и ОГРН работодателя — обязательные поля формы уведомления МВД №8 (приказ №536),
# но в системе их не было. Оставлены пустыми (env-override): по решению — заполняются вручную
# в готовом документе. Когда появятся — задать через переменные окружения EMPLOYER_OKVED/OGRN.
EMPLOYER_OKVED = os.environ.get("EMPLOYER_OKVED", "")
EMPLOYER_OGRN = os.environ.get("EMPLOYER_OGRN", "")
EMPLOYER_SUBDIVISION = os.environ.get("EMPLOYER_SUBDIVISION", "Обособленное подразделение Мурманск")
WORKPLACE_ADDRESS = os.environ.get(
    "WORKPLACE_ADDRESS",
    "Мурманская обл., г. Мурманск, ул. Сполохи 4А, ООО НОВАТЭК-Усть-Луга",
)
DISTRICT_COEFFICIENT = os.environ.get("DISTRICT_COEFFICIENT", "1,500")
CONTRACT_NUMBER_PREFIX = os.environ.get("CONTRACT_NUMBER_PREFIX", "БК-ПСМ-")

# --- Реквизиты госпошлины (квитанция ПД-4сб, налог). УФК по Мурманской области (УМВД).
# Фиксированные, из платёжных реквизитов УМВД (фото от пользователя, приняты как факт).
# Обе пошлины — один получатель, различаются назначением/КБК/суммой (см. DUTY_KINDS ниже).
DUTY_PAYEE_NAME = os.environ.get(
    "DUTY_PAYEE_NAME",
    "УФК по Мурманской области (УМВД России по Мурманской области, л/сч 04491137920)",
)
DUTY_PAYEE_INN = os.environ.get("DUTY_PAYEE_INN", "5191501766")
DUTY_PAYEE_KPP = os.environ.get("DUTY_PAYEE_KPP", "519001001")
DUTY_OKTMO = os.environ.get("DUTY_OKTMO", "47536000")
DUTY_ACCOUNT = os.environ.get("DUTY_ACCOUNT", "03100643000000014900")
DUTY_BANK_NAME = os.environ.get(
    "DUTY_BANK_NAME",
    "ОКЦ № 3 Северо-Западного ГУ Банка России//УФК по Мурманской области г. Мурманск",
)
DUTY_BIC = os.environ.get("DUTY_BIC", "014705901")
# Единый казначейский счёт (ЕКС / корр. счёт для QR по ГОСТ Р 56042). Определяется по БИК,
# подтверждён по официальным источникам (УФК/госорганы Мурманской обл.). Для QR обязателен.
DUTY_CORRESP_ACC = os.environ.get("DUTY_CORRESP_ACC", "40102810745370000041")
# Наименование получателя для QR (без л/сч — в QR идёт чистое наименование).
DUTY_PAYEE_NAME_QR = os.environ.get(
    "DUTY_PAYEE_NAME_QR", "УФК по Мурманской области (УМВД России по Мурманской области)"
)

# Типы пошлин: назначение, КБК (20 цифр, без пробелов), сумма. Ключ передаётся в генератор.
DUTY_KINDS = {
    "registration": {
        "purpose": (
            "Государственная пошлина на постановку иностранного гражданина или лица без "
            "гражданства на учёт по месту пребывания"
        ),
        "kbk": "18810806000010039110",
        "amount": "500",
    },
    "renewal": {
        "purpose": (
            "Государственная пошлина за продление срока временного пребывания иностранного "
            "гражданина в Российской Федерации"
        ),
        "kbk": "18810806000010041110",
        "amount": "1000",
    },
}

MEDICAL_SERVICE_TEXT = (
    "Медицинское освидетельствование на наличие или отсутствие инфекционных заболеваний, "
    "представляющих опасность для окружающих и являющихся основанием для отказа в выдаче "
    "либо аннулирования разрешения на временное проживание иностранных лиц и лиц без "
    "гражданства, или вида на жительство, или патента, или разрешения на работу в Российской "
    "Федерации, если иное не предусмотрено международным договором Российской Федерации, "
    "с проведением лабораторных исследований, проведением осмотра врачом-дерматовенерологом, "
    "осмотра врачом-инфекционистом, с выдачей медицинского заключения и сертификата об "
    "отсутствии у иностранного гражданина заболевания, вызываемого вирусом иммунодефицита "
    "человека (ВИЧ-инфекции)."
)


def _contract_header_parts() -> tuple[str, str, str]:
    """Разбивает CLINIC_CONTRACT_DATE (ожидается формат ДД.ММ.ГГГГ) на день/месяц/год —
    в бумажном бланке это три отдельных поля («25» 06 2026г.), не одна строка.
    [Предполагаю] формат хранения даты в env var — ДД.ММ.ГГГГ, как и в остальных местах
    проекта. Если строку не удалось разобрать — используем как есть одним куском в поле
    "день", чтобы не потерять данные и не упасть, но это не будет визуально соответствовать
    трём полям бланка."""
    try:
        parsed = datetime.strptime(CLINIC_CONTRACT_DATE, "%d.%m.%Y").date()
        return parsed.strftime("%d"), parsed.strftime("%m"), parsed.strftime("%Y")
    except ValueError:
        return CLINIC_CONTRACT_DATE, "", ""


def _format_salary(salary: str) -> str:
    """Оклад из формы (число или строка) -> '30 000' с пробелом-разделителем тысяч.
    Если не парсится как число — возвращаем как есть (кадровик мог ввести строкой)."""
    s = (salary or "").strip().replace(" ", "").replace("\xa0", "")
    if s.isdigit():
        return f"{int(s):,}".replace(",", " ")
    return salary or DASH


def generate_labor_contract_docx(
    employee: Employee,
    position: str,
    salary: str,
    contract_date: date,
    output_dir: str = "/tmp",
) -> str:
    """Трудовой договор с ООО «ТРЕСТСТРОЙМОНТАЖ».

    Переменные из карточки: ФИО, дата рождения, паспорт (серия/номер), адрес (SITE_ADDRESS),
    табельный номер (employee.tab_number). Номер договора = CONTRACT_NUMBER_PREFIX + tab_number.
    Из формы генерации: position (должность), salary (оклад), contract_date (дата договора).
    Прочерк в документе: «паспорт кем/когда выдан» (в модели нет).

    ВНИМАНИЕ: без tab_number номер договора будет обрывком (префикс без номера) — вызывающий
    код (webforms) обязан блокировать генерацию при пустом tab_number. Здесь ставим DASH,
    чтобы не молчать, если всё же вызвали без номера.
    Реквизиты ООО и текст договора приняты как факт без юридической проверки."""
    doc = Document()
    _set_default_style(doc)

    tab = (employee.tab_number or "").strip()
    contract_no = f"{CONTRACT_NUMBER_PREFIX}{tab}" if tab else DASH
    date_str = contract_date.strftime("%d.%m.%Y") if contract_date else DASH
    birth = _date_or_dash(employee.birth_date)
    passport = _passport_str(employee)
    salary_fmt = _format_salary(salary)
    position = (position or "").strip() or DASH

    def h(text):
        p = doc.add_paragraph()
        r = p.add_run(text)
        r.bold = True
        return p

    def para(text):
        return doc.add_paragraph(text)

    # Заголовок
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run(f"ТРУДОВОЙ ДОГОВОР № {contract_no}")
    r.bold = True

    # Город слева, дата справа — через таблицу 1×2 без границ (табы съезжают по ширине поля).
    head = doc.add_table(rows=1, cols=2)
    head.autofit = True
    head.cell(0, 0).paragraphs[0].add_run("г. Москва")
    _rp = head.cell(0, 1).paragraphs[0]
    _rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _rp.add_run(date_str)

    para(
        f"{EMPLOYER_NAME_FULL}, именуемое в дальнейшем «Работодатель», в лице "
        f"Генерального директора {EMPLOYER_DIRECTOR_FULL}, действующего на основании Устава, "
        f"с одной стороны, и {employee.full_name}, именуемый в дальнейшем «Работник», "
        f"с другой стороны, заключили настоящий трудовой договор о нижеследующем:"
    )

    h("1. ПРЕДМЕТ И СРОК ДЕЙСТВИЯ ДОГОВОРА")
    para(
        f"1.1. Согласно настоящему договору Работник принимается на работу в "
        f"{EMPLOYER_SUBDIVISION} {EMPLOYER_NAME_SHORT} на должность {position} с {date_str} года."
    )
    para(
        "1.2. Работник обязуется выполнять все работы, обуславливаемые должностью, на которую "
        "он принимается, а также трудовыми обязанностями и конкретными заданиями (поручениями), "
        "устанавливаемыми Работодателем, и должностной инструкцией в случае ее наличия."
    )
    para(f"1.3. Место работы определено: {WORKPLACE_ADDRESS}")
    para("1.4. Работа по настоящему Договору является для Работника основным местом работы.")
    para("1.5. Срок действия настоящего трудового договора устанавливается на неопределенный срок.")

    h("2. УСЛОВИЯ ТРУДА")
    para(
        "2.1. Работнику устанавливается пятидневная рабочая неделя, восьмичасовой рабочий день. "
        "Рабочий день начинается в 9 часов 00 минут утра, если при приеме на работу в связи с "
        "производственной необходимостью не оговорен другой режим рабочего времени."
    )
    para(
        "Продолжительность перерыва для отдыха и питания составляет 1 час (шестьдесят) минут в "
        "день. Время перерыва определяется на усмотрение Работодателя в пределах между 13-00 и "
        "15-00 часами дня."
    )
    para("2.2. Работник имеет право на ежегодный оплачиваемый отпуск продолжительностью 28 календарных дней.")
    para(
        "2.3. Работодатель осуществляет обязательное социальное, медицинское и пенсионное "
        "страхования Работника в порядке, определенном действующим законодательством."
    )
    para(
        "2.4. Работодатель выплачивает Работнику пособие по временной нетрудоспособности в "
        "размере, установленном действующим законодательством РФ."
    )

    h("3. ОПЛАТА ТРУДА")
    para(
        "3.1. Согласно настоящему договору Работнику выплачивается заработная плата в "
        "соответствии со штатным расписанием. На момент заключения договора заработная плата "
        f"состоит из: Оклад: {salary_fmt} руб.; Районный коэффициент: {DISTRICT_COEFFICIENT}."
    )
    para(
        "3.2. Заработная плата выплачивается Работнику не реже чем каждые полмесяца путем "
        "перечисления денежных средств на банковский счёт Работника, реквизиты которого "
        "Работник сообщает Работодателю в письменном виде."
    )

    h("4. ПРАВА И ОБЯЗАННОСТИ СТОРОН")
    para("4.1. Работник имеет права и обязуется исполнять обязанности, предусмотренные статьей 21 ТК РФ.")
    para("4.2. Работодатель имеет права и обязуется исполнять обязанности, предусмотренные статьей 22 ТК РФ.")

    h("5. КОНФИДЕНЦИАЛЬНОСТЬ")
    para(
        "5.1. Работник обязан обеспечить сохранность, не разглашать и не передавать третьим лицам "
        "сведения и документы, составляющие служебную, коммерческую, техническую, технологическую "
        "или экономическую тайну Работодателя и его клиентов (заказчиков) как в течение срока "
        "действия настоящего Договора, так и в течение 3 лет после его прекращения."
    )

    h("6. ПОРЯДОК УРЕГУЛИРОВАНИЯ СПОРОВ")
    para(
        "6.1. Все споры и разногласия, которые могут возникнуть из настоящего трудового договора "
        "или в связи с ним, будут по возможности решаться сторонами путем переговоров. В случае "
        "недостижения согласия спор подлежит урегулированию в порядке, предусмотренном трудовым "
        "законодательством РФ."
    )

    h("7. ИНЫЕ УСЛОВИЯ")
    para(
        "7.1. Во всем остальном, что не предусмотрено настоящим трудовым договором, стороны "
        "руководствуются законодательством РФ, регулирующим трудовые отношения."
    )
    para("7.2. Настоящий договор составлен в 2 экземплярах, по одному для каждой из сторон.")

    h("8. АДРЕСА СТОРОН И ПОДПИСИ")
    # Две стороны рядом: слева Работник, справа Работодатель. Таблица 1x2 без границ.
    sig = doc.add_table(rows=1, cols=2)
    sig.autofit = True

    left = sig.cell(0, 0)
    left.paragraphs[0].add_run("Работник:").bold = True
    left.add_paragraph(f"{employee.full_name}, {birth} года рождения")
    left.add_paragraph(f"Паспорт: {passport}, выдан: {DASH}")
    left.add_paragraph(f"Адрес: {SITE_ADDRESS}")
    left.add_paragraph("")
    left.add_paragraph("Подпись: _______________")

    right = sig.cell(0, 1)
    right.paragraphs[0].add_run("Работодатель:").bold = True
    right.add_paragraph(EMPLOYER_NAME_FULL)
    right.add_paragraph(f"ИНН: {EMPLOYER_INN} КПП: {EMPLOYER_KPP}")
    right.add_paragraph(f"Юридический адрес: {EMPLOYER_LEGAL_ADDRESS}")
    right.add_paragraph(f"Фактический адрес: {EMPLOYER_ACTUAL_ADDRESS}")
    right.add_paragraph(f"Телефон: {EMPLOYER_PHONE}")
    right.add_paragraph("")
    right.add_paragraph(f"Ген. директор _______________")
    right.add_paragraph(EMPLOYER_DIRECTOR_SHORT)
    right.add_paragraph("м.п.")

    safe_tab = tab or "no_tab"
    filename = f"labor_contract_{employee.id}_{safe_tab}.docx"
    path = os.path.join(output_dir, filename)
    doc.save(path)
    return path


def _build_duty_qr_payload(kind: str) -> str:
    """Платёжная строка по ГОСТ Р 56042 для банковского QR. Порядок и состав полей —
    обязательные (Name, PersonalAcc, BankName, BIC, CorrespAcc, PayeeINN) плюс KPP, KBK,
    OKTMO, Sum (в копейках), Purpose. Sum: рубли*100."""
    spec = DUTY_KINDS[kind]
    fields = {
        "Name": DUTY_PAYEE_NAME_QR,
        "PersonalAcc": DUTY_ACCOUNT,
        "BankName": DUTY_BANK_NAME,
        "BIC": DUTY_BIC,
        "CorrespAcc": DUTY_CORRESP_ACC,
        "PayeeINN": DUTY_PAYEE_INN,
        "KPP": DUTY_PAYEE_KPP,
        "KBK": spec["kbk"],
        "OKTMO": DUTY_OKTMO,
        "Sum": str(int(spec["amount"]) * 100),
        "Purpose": spec["purpose"],
    }
    return "ST00012|" + "|".join(f"{k}={v}" for k, v in fields.items())


def _generate_duty_qr_png(kind: str, out_path: str) -> bool:
    """Строит QR-картинку платёжной строки в out_path. Возвращает True при успехе.
    Библиотека qrcode может отсутствовать (офлайн-сборка) — тогда возвращаем False,
    и квитанция генерируется без QR (не роняем весь документ из-за отсутствия картинки).
    На Railway qrcode должна быть в зависимостях (добавить 'qrcode[pil]' в requirements)."""
    try:
        import qrcode
    except ImportError:
        return False
    try:
        img = qrcode.make(_build_duty_qr_payload(kind))
        img.save(out_path)
        return True
    except Exception:
        return False


def _fix_table_width(table, widths_mm):
    """Жёстко фиксирует ширину колонок таблицы через XML, иначе Word/LibreOffice игнорируют
    cell.width и растягивают таблицу на всю ширину листа. Ставит tblLayout=fixed, общую
    ширину таблицы и tcW (в twips) на каждую ячейку. widths_mm — список ширин колонок в мм."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    def _twips(mm):
        return str(int(mm * 56.6929))  # 1 мм = 56.6929 twips

    tbl = table._tbl
    tblPr = tbl.tblPr

    # Инлайн-границы (не через стиль): мобильные docx-читалки игнорируют границы стиля
    # Table Grid, поэтому прописываем tblBorders прямо в XML — видно в любой программе.
    borders = tblPr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tblPr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = borders.find(qn(f"w:{edge}"))
        if el is None:
            el = OxmlElement(f"w:{edge}")
            borders.append(el)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")  # 0.5pt
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "000000")

    # layout = fixed
    layout = tblPr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tblPr.append(layout)
    layout.set(qn("w:type"), "fixed")

    # общая ширина таблицы
    total = sum(widths_mm)
    tblW = tblPr.find(qn("w:tblW"))
    if tblW is None:
        tblW = OxmlElement("w:tblW")
        tblPr.append(tblW)
    tblW.set(qn("w:w"), _twips(total))
    tblW.set(qn("w:type"), "dxa")

    # tblGrid — ширины колонок
    grid = tbl.find(qn("w:tblGrid"))
    if grid is not None:
        for gc, w in zip(grid.findall(qn("w:gridCol")), widths_mm):
            gc.set(qn("w:w"), _twips(w))

    # tcW на каждую ячейку
    for row in table.rows:
        for cell, w in zip(row.cells, widths_mm):
            tcPr = cell._tc.get_or_add_tcPr()
            tcW = tcPr.find(qn("w:tcW"))
            if tcW is None:
                tcW = OxmlElement("w:tcW")
                tcPr.append(tcW)
            tcW.set(qn("w:w"), _twips(w))
            tcW.set(qn("w:type"), "dxa")


def generate_duty_receipt_docx(kind: str, employee=None, payer_name: str = "", output_dir: str = "/tmp") -> str:
    """Квитанция на оплату госпошлины, форма ПД-4сб(налог). kind: "registration" (500, постановка
    на учёт) или "renewal" (1000, продление пребывания). Реквизиты фиксированные (DUTY_*), сумма
    фиксированная по типу. Плательщик (ФИО/адрес) — поля предусмотрены, но пустые: заполняются
    вручную (пользователь: ФИО пока не подставляем). employee — задел на будущее автозаполнение.

    ПД-4сб состоит из двух одинаковых частей: «Извещение» (остаётся в банке) и «Квитанция»
    (у плательщика). Обе части идентичны — строятся одной вспомогательной функцией."""
    if kind not in DUTY_KINDS:
        raise ValueError(f"Неизвестный тип пошлины: {kind!r}. Ожидается один из {list(DUTY_KINDS)}")
    spec = DUTY_KINDS[kind]

    doc = Document()
    _set_default_style(doc)

    # Компоновка на один лист A4: узкие поля + мелкий шрифт таблицы. Две части ПД-4сб (13 строк
    # каждая) + QR иначе не влезают. Шрифт таблицы 8pt — компромисс ради одного листа (запрошено).
    from docx.shared import Pt as _Pt, Inches as _In
    for sec in doc.sections:
        sec.top_margin = _In(0.5)
        sec.bottom_margin = _In(0.5)
        sec.left_margin = _In(0.6)
        sec.right_margin = _In(0.6)

    # ФИО плательщика — из формы (кадровик вводит; плательщик НЕ обязательно работник).
    payer_fio = (payer_name or "").strip() or DASH
    payer_addr = DASH  # адрес плательщика — прочерк (по решению: не вводим)
    # ФИО работника, за кого платёж, добавляется в назначение платежа.
    _worker = (getattr(employee, "full_name", "") or "").strip() if employee else ""

    # QR один раз генерируем во временный файл, вставляем в обе части (извещение и квитанция).
    import tempfile
    qr_path = os.path.join(tempfile.gettempdir(), f"duty_qr_{kind}.png")
    qr_ok = _generate_duty_qr_png(kind, qr_path)

    def _tiny(paragraph, size=10):
        for r in paragraph.runs:
            r.font.size = _Pt(size)
        # убрать вертикальные пробелы между строками таблицы
        pf = paragraph.paragraph_format
        pf.space_before = _Pt(1)
        pf.space_after = _Pt(1)
        pf.line_spacing = 1.0

    def _build_part(part_title: str) -> None:
        head = doc.add_paragraph()
        head.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Заголовок части держится со следующим содержимым — часть не начинается в самом низу листа
        # с переносом таблицы/QR на новую страницу.
        head.paragraph_format.keep_with_next = True
        head.paragraph_format.space_before = _Pt(2)
        head.paragraph_format.space_after = _Pt(1)
        r = head.add_run(part_title)
        r.bold = True
        r.font.size = _Pt(13)
        sub = doc.add_paragraph()
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sr = sub.add_run("Форма № ПД-4сб (налог)")
        sr.italic = True
        sr.font.size = _Pt(9)

        rows = [
            ("Наименование получателя", DUTY_PAYEE_NAME),
            ("ИНН / КПП получателя", f"{DUTY_PAYEE_INN} / {DUTY_PAYEE_KPP}"),
            ("Код ОКТМО", DUTY_OKTMO),
            ("Счёт получателя платежа", DUTY_ACCOUNT),
            ("Банк получателя", DUTY_BANK_NAME),
            ("БИК / Кор. счёт", f"{DUTY_BIC} / {DUTY_CORRESP_ACC}"),
            ("КБК", spec["kbk"]),
            ("Наименование платежа", spec["purpose"] + (f" за {_worker}" if _worker else "")),
            ("Ф.И.О. плательщика", payer_fio),
            ("Адрес плательщика", payer_addr),
            ("Сумма платежа", f"{spec['amount']} руб. 00 коп."),
        ]
        table = doc.add_table(rows=len(rows), cols=2)
        table.style = "Table Grid"
        table.autofit = False
        table.allow_autofit = False
        for i, (label, value) in enumerate(rows):
            c0 = table.cell(i, 0).paragraphs[0]
            run = c0.add_run(label)
            run.bold = True
            _tiny(c0)
            c1 = table.cell(i, 1).paragraphs[0]
            c1.add_run(value)
            _tiny(c1)
        # Жёсткая фиксация ширины через XML: подписи 55 мм, значения 110 мм (итого 165 мм —
        # ширина печатного поля A4 при полях 0.5"/0.6"). Без этого таблица растягивается.
        _fix_table_width(table, [50, 95])

        # QR + подпись рядом: таблица 1x2 без границ. Слева QR, справа поля подписи.
        foot = doc.add_table(rows=1, cols=2)
        foot.autofit = True
        if qr_ok:
            foot.cell(0, 0).paragraphs[0].add_run().add_picture(qr_path, width=_In(1.1))
        else:
            foot.cell(0, 0).paragraphs[0].add_run("(QR недоступен)").font.size = _Pt(7)
        rc = foot.cell(0, 1)
        p1 = rc.paragraphs[0]; p1.add_run("Плательщик (подпись) _______________"); _tiny(p1)
        p1.paragraph_format.keep_together = True
        p2 = rc.add_paragraph(); p2.add_run("Дата _______________"); _tiny(p2)
        p2.paragraph_format.keep_together = True
        p3 = rc.add_paragraph(); p3.add_run("Отсканируйте QR в банковском приложении для оплаты"); _tiny(p3, 7)

    _build_part("ИЗВЕЩЕНИЕ")
    _sep = doc.add_paragraph("- - - - - - - - - - - - - - - - - - - - - - - - - - - - - -")
    _sep.paragraph_format.space_before = _Pt(2)
    _sep.paragraph_format.space_after = _Pt(2)
    for _r in _sep.runs:
        _r.font.size = _Pt(8)
    _build_part("КВИТАНЦИЯ")

    filename = f"duty_receipt_{kind}.docx"
    path = os.path.join(output_dir, filename)
    doc.save(path)
    return path


def generate_termination_notice_docx(employee: Employee, basis: str = "",
                                     output_dir: str = "/tmp") -> str:
    """Уведомление МВД о прекращении (расторжении) трудового договора с иностранцем.
    Структура по форме №8 приказа МВД России от 30.07.2020 №536 (действует до 01.09.2026 —
    с этой даты приказ №290 от 12.05.2026 вводит новые формы; см. задачу на обновление).

    НЕ пиксельная копия бланка: подача основная через Госуслуги (там своя форма), это документ
    для подготовки данных и архива / запасной бумажной подачи. Обязательные поля, которых нет в
    системе (ОКВЭД, ОГРН работодателя), — прочерком под ручное заполнение.

    basis — основание расторжения (по собственному желанию / инициатива работодателя /
    истечение срока и т.п.), вводится в форме увольнения. Дата расторжения — contract_end_date."""
    doc = Document()
    _set_default_style(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("УВЕДОМЛЕНИЕ\nо прекращении (расторжении) трудового договора\n"
                      "с иностранным гражданином")
    r.bold = True
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub.add_run("(форма по приложению №8 к приказу МВД России от 30.07.2020 №536)")
    sr.italic = True
    sr.font.size = Pt(9)

    doc.add_paragraph()
    doc.add_paragraph("В ____________________________________________________________")
    doc.add_paragraph("(наименование территориального органа МВД России)")
    doc.add_paragraph()

    doc.add_paragraph("1. Сведения о работодателе (заказчике работ, услуг):")
    rows = [
        ("Полное наименование", EMPLOYER_NAME_FULL),
        ("ИНН", EMPLOYER_INN),
        ("КПП", EMPLOYER_KPP),
        ("ОГРН", _text_or_dash(EMPLOYER_OGRN)),
        ("ОКВЭД (основной вид деятельности)", _text_or_dash(EMPLOYER_OKVED)),
        ("Юридический адрес", EMPLOYER_LEGAL_ADDRESS),
        ("Фактический адрес (место работы)", WORKPLACE_ADDRESS),
        ("Телефон", EMPLOYER_PHONE),
    ]
    t1 = doc.add_table(rows=len(rows), cols=2)
    t1.style = "Table Grid"
    for i, (k, v) in enumerate(rows):
        t1.cell(i, 0).paragraphs[0].add_run(k).bold = True
        t1.cell(i, 1).paragraphs[0].add_run(v)
    _fix_table_width(t1, [70, 100])

    doc.add_paragraph()
    doc.add_paragraph("2. Сведения об иностранном гражданине (лице без гражданства):")
    name_parts = (employee.full_name or "").split()
    surname = name_parts[0] if name_parts else DASH
    first = name_parts[1] if len(name_parts) > 1 else DASH
    patr = name_parts[2] if len(name_parts) > 2 else DASH
    rows2 = [
        ("Фамилия", surname),
        ("Имя", first),
        ("Отчество", patr),
        ("Гражданство", _text_or_dash(getattr(employee, "citizenship", None) or "Республика Казахстан")),
        ("Дата рождения", _date_or_dash(employee.birth_date)),
        ("Документ, удостоверяющий личность (серия, номер)", _passport_str(employee)),
    ]
    t2 = doc.add_table(rows=len(rows2), cols=2)
    t2.style = "Table Grid"
    for i, (k, v) in enumerate(rows2):
        t2.cell(i, 0).paragraphs[0].add_run(k).bold = True
        t2.cell(i, 1).paragraphs[0].add_run(v)
    _fix_table_width(t2, [70, 100])

    doc.add_paragraph()
    doc.add_paragraph("3. Сведения о трудовом (гражданско-правовом) договоре:")
    _contract_no = (CONTRACT_NUMBER_PREFIX + employee.tab_number.strip()) if (employee.tab_number or "").strip() else DASH
    rows3 = [
        ("Номер договора", _contract_no),
        ("Дата заключения договора", _date_or_dash(employee.contract_date)),
        ("Дата прекращения (расторжения)", _date_or_dash(employee.contract_end_date)),
        ("Основание прекращения (расторжения)", _text_or_dash(basis)),
    ]
    t3 = doc.add_table(rows=len(rows3), cols=2)
    t3.style = "Table Grid"
    for i, (k, v) in enumerate(rows3):
        t3.cell(i, 0).paragraphs[0].add_run(k).bold = True
        t3.cell(i, 1).paragraphs[0].add_run(v)
    _fix_table_width(t3, [70, 100])

    doc.add_paragraph()
    doc.add_paragraph("Уведомление подаётся в срок не более 3 рабочих дней с даты прекращения "
                      "(расторжения) договора (п. 8 ст. 13 Федерального закона от 25.07.2002 №115-ФЗ).")
    doc.add_paragraph()
    doc.add_paragraph(f"Руководитель: _______________ / {EMPLOYER_DIRECTOR_SHORT}")
    doc.add_paragraph(f"Дата подачи: «___» __________ 20__ г.        М.П.")

    filename = f"termination_notice_{employee.id}.docx"
    path = os.path.join(output_dir, filename)
    doc.save(path)
    return path


def generate_departure_notice_docx(employee: Employee, output_dir: str = "/tmp") -> str:
    """Уведомление об убытии иностранного гражданина из места пребывания (снятие с
    миграционного учёта). Принимающая сторона — ООО «ТРЕСТСТРОЙМОНТАЖ» (вахта, предоставляет
    жильё). Срок подачи — 7 рабочих дней с даты убытия (ст. 23 №109-ФЗ, п. 45 Правил ПП №9).

    Дата убытия = contract_end_date (дата увольнения). Адрес пребывания, откуда убыл, —
    SITE_ADDRESS (площадка). НЕ пиксельная копия бланка (подача через Госуслуги / бумага —
    запасной путь); структурированный документ со всеми обязательными сведениями."""
    doc = Document()
    _set_default_style(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("УВЕДОМЛЕНИЕ\nоб убытии иностранного гражданина\nиз места пребывания")
    r.bold = True
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub.add_run("(снятие с миграционного учёта по месту пребывания)")
    sr.italic = True
    sr.font.size = Pt(9)

    doc.add_paragraph()
    doc.add_paragraph("В ____________________________________________________________")
    doc.add_paragraph("(наименование территориального органа МВД России)")
    doc.add_paragraph()

    doc.add_paragraph("1. Сведения об иностранном гражданине, подлежащем снятию с учёта:")
    name_parts = (employee.full_name or "").split()
    surname = name_parts[0] if name_parts else DASH
    first = name_parts[1] if len(name_parts) > 1 else DASH
    patr = name_parts[2] if len(name_parts) > 2 else DASH
    rows = [
        ("Фамилия", surname),
        ("Имя", first),
        ("Отчество", patr),
        ("Гражданство", _text_or_dash(getattr(employee, "citizenship", None) or "Республика Казахстан")),
        ("Дата рождения", _date_or_dash(employee.birth_date)),
        ("Документ, удостоверяющий личность", _passport_str(employee)),
        ("Адрес места пребывания (откуда убыл)", SITE_ADDRESS),
        ("Дата убытия из места пребывания", _date_or_dash(employee.contract_end_date)),
    ]
    t1 = doc.add_table(rows=len(rows), cols=2)
    t1.style = "Table Grid"
    for i, (k, v) in enumerate(rows):
        t1.cell(i, 0).paragraphs[0].add_run(k).bold = True
        t1.cell(i, 1).paragraphs[0].add_run(v)
    _fix_table_width(t1, [70, 100])

    doc.add_paragraph()
    doc.add_paragraph("2. Сведения о принимающей стороне (организация):")
    rows2 = [
        ("Полное наименование", EMPLOYER_NAME_FULL),
        ("ИНН", EMPLOYER_INN),
        ("КПП", EMPLOYER_KPP),
        ("ОГРН", _text_or_dash(EMPLOYER_OGRN)),
        ("Адрес", EMPLOYER_LEGAL_ADDRESS),
        ("Телефон", EMPLOYER_PHONE),
        ("Ф.И.О. представителя", DASH),
        ("Документ, удостоверяющий личность представителя", DASH),
    ]
    t2 = doc.add_table(rows=len(rows2), cols=2)
    t2.style = "Table Grid"
    for i, (k, v) in enumerate(rows2):
        t2.cell(i, 0).paragraphs[0].add_run(k).bold = True
        t2.cell(i, 1).paragraphs[0].add_run(v)
    _fix_table_width(t2, [70, 100])

    doc.add_paragraph()
    doc.add_paragraph("Уведомление об убытии представляется принимающей стороной в срок 7 рабочих "
                      "дней с даты убытия (ст. 23 Федерального закона от 18.07.2006 №109-ФЗ, "
                      "п. 45 Правил, утв. постановлением Правительства РФ от 15.01.2007 №9).")
    doc.add_paragraph()
    doc.add_paragraph("Представитель принимающей стороны: _______________ / _______________")
    doc.add_paragraph("Дата подачи: «___» __________ 20__ г.        М.П.")

    filename = f"departure_notice_{employee.id}.docx"
    path = os.path.join(output_dir, filename)
    doc.save(path)
    return path


def generate_medical_referral_docx(employee: Employee, output_dir: str = "/tmp") -> str:
    """Направление на медицинское освидетельствование — форма ГОАУЗ «МОМЦ» (Приложение №1
    к договору), заполняется по факту согласования конкретной даты/кабинета с клиникой.

    Дата приёма, номер кабинета и время намеренно оставлены пустыми полями для ручного
    заполнения — это отдельный процесс согласования с клиникой, бот не может знать
    расписание клиники заранее и не должен его придумывать."""
    missing = _require_fields(employee, MEDICAL_REFERRAL_REQUIRED_FIELDS)

    doc = Document()
    _set_default_style(doc)

    if missing:
        _add_test_warning_paragraph(doc, missing)

    day, month, year = _contract_header_parts()
    header = doc.add_paragraph()
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header.add_run(
        f"к Договору № {CLINIC_CONTRACT_NUMBER} от «{day}» {month} {year}г.\nПриложение № 1"
    )

    doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("НАПРАВЛЕНИЕ НА МЕДИЦИНСКОЕ ОСВИДЕТЕЛЬСТВОВАНИЕ")
    run.bold = True
    run.font.size = Pt(13)

    doc.add_paragraph()
    doc.add_paragraph(f"В {CLINIC_NAME}")
    doc.add_paragraph("наименование медицинской организации (МО)")

    name_parts = employee.full_name.split()
    surname = name_parts[0] if name_parts else DASH
    first_name = name_parts[1] if len(name_parts) > 1 else DASH
    patronymic = name_parts[2] if len(name_parts) > 2 else DASH

    doc.add_paragraph(f"1. Фамилия {surname}")
    doc.add_paragraph(f"Имя {first_name}")
    doc.add_paragraph(f"Отчество {patronymic}")

    birth = _date_or_dash(employee.birth_date)
    doc.add_paragraph(f"2. Дата рождения (число, месяц, год) {birth}")

    doc.add_paragraph(f"3. Адрес (по месту проживания) {SITE_ADDRESS}")

    doc.add_paragraph(
        f"4. Серия паспорта {_text_or_dash(employee.passport_series)}  "
        f"Номер паспорта {_text_or_dash(employee.passport_number)}"
    )

    doc.add_paragraph(f"5. Место работы {PAYER_NAME}")

    doc.add_paragraph("6. Наименование медицинской услуги (медицинского освидетельствования)")
    doc.add_paragraph(MEDICAL_SERVICE_TEXT)

    doc.add_paragraph("7. Дата проведения услуги _____________ кабинет N _____ время _____")

    doc.add_paragraph(
        "8. Полное наименование организации, направившей иностранного гражданина, "
        f"телефон {PAYER_PHONE}"
    )
    doc.add_paragraph(PAYER_NAME)
    doc.add_paragraph("подпись, печать _____________________")

    doc.add_paragraph(f"10. Дата выдачи направления {datetime.now(MSK).date().strftime('%d.%m.%Y')}")

    doc.add_paragraph()

    # Блок подписей "От Исполнителя / От Заказчика" — был в бумажном бланке, отсутствовал
    # в генераторе. Таблица 2×2 без границ, чтобы визуально повторить два столбца оригинала.
    table = doc.add_table(rows=4, cols=2)
    table.autofit = True

    def _cell(row: int, col: int, text: str, bold: bool = False) -> None:
        cell = table.cell(row, col)
        p = cell.paragraphs[0]
        r = p.add_run(text)
        r.bold = bold

    _cell(0, 0, "От Исполнителя:", bold=True)
    _cell(0, 1, "От Заказчика:", bold=True)
    _cell(1, 0, f"{CLINIC_SHORT_NAME}")
    _cell(1, 1, "Индивидуальный предприниматель")
    _cell(2, 0, "Главный врач")
    _cell(2, 1, "")
    _cell(3, 0, f"_____________________ {CLINIC_CHIEF_DOCTOR_NAME}\nм.п.")
    _cell(3, 1, f"_____________________ {PAYER_SIGNATORY_NAME}\nм.п.")

    filename = f"medical_referral_{employee.id}.docx"
    path = os.path.join(output_dir, filename)
    doc.save(path)
    return path


def generate_employees_xlsx(employees: list[Employee], output_dir: str = "/tmp") -> str:
    """Полный список сотрудников таблицей — замена постраничному тексту в /employees.
    Отдельный файл от Google Sheets: тот обновляется по крону раз в сутки командой
    export_to_sheets_api.py, этот — генерируется по запросу с текущим состоянием БД."""
    wb = Workbook()
    ws = wb.active
    ws.title = "Сотрудники"

    headers = [
        "ФИО", "Гражданство", "Категория", "Дата въезда", "Дата договора",
        "Статус занятости", "Согласие", "Телефон", "Язык",
        "Дата рождения", "Серия паспорта", "Номер паспорта", "Адрес", "Откуда въехал",
    ]
    ws.append(headers)
    for cell in ws[1]:
        cell.font = Font(bold=True)

    for emp in employees:
        ws.append([
            emp.full_name,
            emp.citizenship or "",
            emp.category.value if emp.category else "",
            emp.entry_date.strftime("%d.%m.%Y") if emp.entry_date else "",
            emp.contract_date.strftime("%d.%m.%Y") if emp.contract_date else "",
            emp.employment_status or "",
            "да" if emp.consent_status.value == "confirmed" else "нет",
            emp.phone or "",
            emp.language or "",
            emp.birth_date.strftime("%d.%m.%Y") if emp.birth_date else "",
            emp.passport_series or "",
            emp.passport_number or "",
            emp.address or "",
            emp.entry_country or "",
        ])

    for col in ws.columns:
        max_len = max((len(str(c.value)) for c in col if c.value), default=10)
        ws.column_dimensions[col[0].column_letter].width = min(max_len + 2, 40)

    path = os.path.join(output_dir, "employees.xlsx")
    wb.save(path)
    return path
